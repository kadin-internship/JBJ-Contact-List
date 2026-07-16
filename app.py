import os
import io
import re
import csv
from datetime import date, datetime, timedelta
from dotenv import load_dotenv
from flask import Flask, request, jsonify, send_file, render_template, render_template_string, redirect, url_for, abort
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from config import Config
from db import db
from models import Contact, OutreachOrg, Activity, User, AuditLog, CaseStudy, EmailTemplate, FlyerTemplate, FlyerAsset, Task, SocialToken, EmailEvent, AvailabilityRule, Booking, LandingPage, LandingPageSubmission, EmailSequence, EmailSequenceStep, EmailSequenceEnrollment, Proposal
from schemas import ContactSchema
from utils import (
    read_uploaded_file, clean_dataframe, clean_outreach_orgs,
    looks_like_contacts_sheet, looks_like_orgs_sheet,
)
from sqlalchemy import or_, and_, func

load_dotenv()

# Error monitoring is opt-in: only initializes if SENTRY_DSN is set, so the
# app runs exactly as before for local dev or anyone who hasn't created a
# Sentry account. Without this, the only way to learn about a production
# error is a user reporting it.
_sentry_dsn = os.environ.get('SENTRY_DSN')
if _sentry_dsn:
    import sentry_sdk
    from sentry_sdk.integrations.flask import FlaskIntegration
    sentry_sdk.init(
        dsn=_sentry_dsn,
        integrations=[FlaskIntegration()],
        traces_sample_rate=0.0,  # error tracking only, not performance tracing
        send_default_pii=False,  # this app holds contact PII -- don't forward any of it
    )

login_manager = LoginManager()
login_manager.login_view = 'login'

# In-memory storage -- fine for brute-force protection on a single login
# form, but each gunicorn worker process counts independently (no shared
# Redis), so the effective limit across the whole fleet is up to
# (per-worker limit x worker count), not a hard global ceiling. Good
# enough to stop naive password guessing; revisit with a shared store
# (e.g. Redis) if that gap ever matters.
limiter = Limiter(key_func=get_remote_address, default_limits=[])


def parse_multi_param(name):
    """Several filters (county, tag, org_tag) are multi-select -- the
    frontend sends the chosen values as a single comma-joined query param
    (e.g. county=Dallas,Tarrant or tag=Chamber,Clergy)."""
    raw = request.args.get(name, type=str)
    if not raw:
        return []
    return [v.strip() for v in raw.split(',') if v.strip()]


def split_multi(value):
    """Normalizes a tag/org_tag/county value that may arrive as a list, a
    comma-joined string, or a single plain string (e.g. from a JSON body)
    into a list."""
    if not value:
        return []
    if isinstance(value, (list, tuple)):
        return [str(v).strip() for v in value if v and str(v).strip()]
    return [v.strip() for v in str(value).split(',') if v.strip()]


def county_filter_clause(counties):
    """Matches if ANY selected county appears as one of the comma-separated
    entries in Contact.county, not just an exact whole-field match --
    e.g. selecting "Dallas" should also match a contact filed under
    "Dallas, Tarrant"."""
    if not counties:
        return None
    clauses = []
    for c in counties:
        clauses.append(or_(
            Contact.county.ilike(c),
            Contact.county.ilike(f"{c}, %"),
            Contact.county.ilike(f"%, {c}"),
            Contact.county.ilike(f"%, {c}, %"),
        ))
    return or_(*clauses)


def contact_incomplete_clause():
    """A contact counts as incomplete when there's no way to reach them --
    no email AND no phone of either kind. Mirrors the "Incomplete"/
    "Complete" flag shown on the contact detail panel (see `incomplete` in
    showContactDetail() in app.js), so the dashboard stat and the per-
    contact flag never disagree. Contact.data_complete is unused here --
    nothing in the UI ever sets it, so it was permanently 0%."""
    no_email = or_(Contact.email.is_(None), Contact.email == '')
    no_phone = and_(
        or_(Contact.phone_office.is_(None), Contact.phone_office == ''),
        or_(Contact.phone_cell.is_(None), Contact.phone_cell == ''),
    )
    return and_(no_email, no_phone)


def filtered_contacts_query(q=None, tag=None, county=None, contact_id=None, org_tag=None, followup=None, favorites_only=False):
    """Shared filter logic for /api/contacts and the export endpoints, so
    exports always match what's currently shown on screen.

    `tag` filters Contact.tag (People view categories). `org_tag` filters by
    organizations whose OutreachOrg.tag matches (Organizations view categories)
    -- the two category sets use different names for the same kind of group,
    so org_tag is resolved to organization names first rather than treated as
    a Contact.tag value.

    `followup` is 'never' (no Activity at all) or a number of days as a
    string ('30'/'60'/'90') meaning "no Activity logged in that many days,
    including never" -- a contact with no activity at all always counts as
    overdue, regardless of the threshold.
    """
    query = Contact.query
    if contact_id:
        query = query.filter(Contact.id == contact_id)
    if q:
        like = f"%{q}%"
        full_name = func.coalesce(Contact.first_name, '') + ' ' + func.coalesce(Contact.last_name, '')
        query = query.filter(or_(
            full_name.ilike(like),
            Contact.organization.ilike(like),
            Contact.title.ilike(like),
            Contact.email.ilike(like),
            Contact.county.ilike(like),
        ))
    if tag:
        tags = split_multi(tag)
        if tags:
            query = query.filter(Contact.tag.in_(tags))
    if org_tag:
        org_tags = split_multi(org_tag)
        org_names = [o[0].lower() for o in db.session.query(OutreachOrg.organization).filter(OutreachOrg.tag.in_(org_tags)).all() if o[0]]
        if not org_names:
            return query.filter(False)
        query = query.filter(func.lower(Contact.organization).in_(org_names))
    if county:
        counties = split_multi(county)
        clause = county_filter_clause(counties)
        if clause is not None:
            query = query.filter(clause)
    if followup:
        last_contacted = (
            db.session.query(Activity.contact_id, func.max(Activity.contacted_on).label('last_contacted'))
            .group_by(Activity.contact_id)
            .subquery()
        )
        query = query.outerjoin(last_contacted, last_contacted.c.contact_id == Contact.id)
        if followup == 'never':
            query = query.filter(last_contacted.c.last_contacted.is_(None))
        else:
            try:
                cutoff = date.today() - timedelta(days=int(followup))
            except (TypeError, ValueError):
                cutoff = None
            if cutoff is not None:
                query = query.filter(or_(
                    last_contacted.c.last_contacted.is_(None),
                    last_contacted.c.last_contacted < cutoff,
                ))
    if favorites_only:
        query = query.filter(Contact.is_favorite == True)
    return query


def log_audit(action, entity_type, entity_id=None, entity_label=None, details=None):
    """Records who did what, for the admin-only Audit Log page. Commits on
    its own -- callers should already have committed the actual change, so
    a failure here never rolls back the change it's describing."""
    entry = AuditLog(
        user_id=current_user.id if current_user.is_authenticated else None,
        actor_name=current_user.display_name if current_user.is_authenticated else 'System',
        action=action,
        entity_type=entity_type,
        entity_id=entity_id,
        entity_label=entity_label,
        details=details,
    )
    db.session.add(entry)
    db.session.commit()


ACTION_LABELS = {
    'contact_created': 'Added contact',
    'contact_updated': 'Edited contact',
    'contact_deleted': 'Deleted contact',
    'spreadsheet_sync': 'Synced spreadsheet',
    'user_created': 'Created user login',
    'password_reset': 'Reset password for',
    'case_study_created': 'Added case study',
    'case_study_updated': 'Edited case study',
    'case_study_deleted': 'Deleted case study',
    'case_study_uploaded': 'Uploaded case study file(s)',
    'email_template_created': 'Created email template',
    'email_template_updated': 'Edited email template',
    'email_template_deleted': 'Deleted email template',
    'email_template_sent': 'Sent email',
    'flyer_template_created': 'Created flyer',
    'flyer_template_updated': 'Edited flyer',
    'flyer_template_deleted': 'Deleted flyer',
    'flyer_template_sent': 'Sent flyer',
}


def format_audit_details(action, details):
    """Turns an AuditLog row's raw `details` JSON into a one-line, readable
    summary for the audit log page -- the shape of `details` differs per
    action (a field-diff for edits, import counts for a sync, etc)."""
    d = details or {}
    if action == 'contact_updated':
        parts = []
        for field, change in d.items():
            parts.append(f"{field}: \"{change.get('old') or ''}\" → \"{change.get('new') or ''}\"")
        return '; '.join(parts)
    if action == 'spreadsheet_sync':
        c = d.get('contacts', {})
        o = d.get('organizations', {})
        return (
            f"Contacts: {c.get('inserted', 0)} new, {c.get('updated', 0)} updated · "
            f"Organizations: {o.get('inserted', 0)} new, {o.get('updated', 0)} updated"
        )
    if action == 'user_created':
        return 'Admin access' if d.get('is_admin') else 'Standard access'
    if action == 'case_study_uploaded':
        titles = d.get('titles') or []
        return f"{d.get('count', len(titles))} file(s): {', '.join(titles[:5])}{'...' if len(titles) > 5 else ''}"
    if action == 'email_template_sent':
        attachment_count = d.get('attachment_count', 0)
        suffix = f" with {attachment_count} attachment(s)" if attachment_count else ''
        return f"To: {d.get('to', '')}{suffix}"
    return ''


def format_audit_summary(entry):
    """Builds the one-line "what happened" summary shown per row on the
    Audit Log page, e.g. 'Edited contact Jane Doe -- title: "" -> "Mayor"'."""
    action_label = ACTION_LABELS.get(entry.action, entry.action)
    line = f"{action_label} {entry.entity_label}" if entry.entity_label else action_label
    detail = format_audit_details(entry.action, entry.details)
    return f"{line} - {detail}" if detail else line


def _bootstrap_admin_user():
    """Create one admin account from env vars if no user exists yet.

    Render's Shell tab (the obvious way to run create_user.py against the
    production database) needs a paid plan. This lets the very first
    account get created from env vars set in Render's free Environment
    tab instead. Only fires while the users table is empty, so it can't
    be replayed later to reset someone's password.
    """
    if User.query.count() > 0:
        return
    username = os.environ.get('BOOTSTRAP_ADMIN_USERNAME')
    password = os.environ.get('BOOTSTRAP_ADMIN_PASSWORD')
    if not username or not password:
        return
    display_name = os.environ.get('BOOTSTRAP_ADMIN_DISPLAY_NAME', username)
    user = User(username=username, display_name=display_name, is_admin=True)
    user.set_password(password)
    db.session.add(user)
    db.session.commit()


def _import_contacts(df, result):
    """Upsert a People-sheet DataFrame into Contact, accumulating
    inserted/updated/skipped counts into `result`. Matches existing rows
    by email when present; rows without an email (allowed since
    Contact.email is optional) are matched by name+organization instead,
    since there's no other reliable natural key."""
    cleaned = clean_dataframe(df)
    for row in cleaned:
        email = (row.get('email') or '').strip()
        if email:
            existing = Contact.query.filter(func.lower(Contact.email) == email.lower()).first()
        else:
            existing = Contact.query.filter(
                func.lower(Contact.first_name) == (row.get('first_name') or '').lower(),
                func.lower(Contact.last_name) == (row.get('last_name') or '').lower(),
                func.lower(Contact.organization) == (row.get('organization') or '').lower(),
            ).first()
        if existing:
            changed = False
            for field in ['first_name', 'last_name', 'organization', 'title', 'phone_office', 'phone_cell', 'active', 'county', 'notes', 'tag']:
                val = row.get(field)
                if val and (getattr(existing, field) in (None, '', False)):
                    setattr(existing, field, val)
                    changed = True
            existing_lists = existing.lists or []
            new_lists = row.get('lists') or []
            merged = list(dict.fromkeys(existing_lists + new_lists))
            if merged != existing_lists:
                existing.lists = merged
                changed = True
            if existing.data_complete != bool(row.get('data_complete')):
                existing.data_complete = bool(row.get('data_complete'))
                changed = True
            if changed:
                db.session.add(existing)
                result['updated'] += 1
            else:
                result['skipped'] += 1
        else:
            c = Contact(
                tag=row.get('tag') or None,
                organization=row.get('organization') or None,
                first_name=row.get('first_name') or None,
                last_name=row.get('last_name') or None,
                title=row.get('title') or None,
                phone_office=row.get('phone_office') or None,
                phone_cell=row.get('phone_cell') or None,
                email=email or None,
                active=row.get('active') or None,
                lists=row.get('lists') or [],
                county=row.get('county') or None,
                notes=row.get('notes') or None,
                data_complete=bool(row.get('data_complete')),
            )
            db.session.add(c)
            result['inserted'] += 1


def _import_orgs(df, result):
    """Upsert an Organizations-sheet DataFrame into OutreachOrg, matching
    existing rows by (tag, organization)."""
    cleaned = clean_outreach_orgs(df)
    for row in cleaned:
        existing = OutreachOrg.query.filter_by(tag=row['tag'], organization=row['organization']).first()
        if existing:
            changed = False
            if row.get('updated') and existing.updated != row['updated']:
                existing.updated = row['updated']
                changed = True
            if row.get('notes') and existing.notes != row['notes']:
                existing.notes = row['notes']
                changed = True
            if changed:
                db.session.add(existing)
                result['updated'] += 1
            else:
                result['skipped'] += 1
        else:
            rec = OutreachOrg(tag=row['tag'], organization=row['organization'], updated=row.get('updated'), notes=row.get('notes'))
            db.session.add(rec)
            result['inserted'] += 1


def extract_case_study_text(file_storage):
    """Pulls raw text out of an uploaded .pdf or .docx for the case-study
    importer to hand to Claude. Returns (text, error) -- exactly one is
    None. Native Google Docs aren't readable here (no Drive API access);
    the user has to export them to .docx or PDF first via Drive's
    Download menu, same as old binary .doc files."""
    filename = file_storage.filename or ''
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext == 'pdf':
        from pypdf import PdfReader
        try:
            reader = PdfReader(file_storage.stream)
            text = '\n'.join((page.extract_text() or '') for page in reader.pages)
        except Exception as e:
            return None, f'Could not read this PDF: {e}'
    elif ext == 'docx':
        from docx import Document as DocxReader
        try:
            doc = DocxReader(file_storage.stream)
            text = '\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            return None, f'Could not read this Word document: {e}'
    else:
        return None, (
            f'Unsupported file type ".{ext or "unknown"}" -- only .pdf and .docx are supported. '
            'Native Google Docs need to be exported first (File > Download > Microsoft Word or PDF).'
        )

    text = text.strip()
    if len(text) < 30:
        return None, 'Could not find readable text in this file (it may be a scanned/image-only PDF).'
    return text, None


def absolutize_static_urls(html, base_url):
    """Rewrites relative /static/... src/href references (e.g. the email
    builder's "Insert logo" button) into absolute URLs using the current
    request's host. A relative path only resolves inside the browser tab
    it was inserted in -- a recipient's email client has no "current
    page" to resolve it against, so it just shows a broken image."""
    base = base_url.rstrip('/')
    return re.sub(r'(src|href)="(/static/[^"]*)"', lambda m: f'{m.group(1)}="{base}{m.group(2)}"', html)


def _sg_from():
    """Returns (from_email, from_name) from env vars."""
    from_email = os.environ.get('SMTP_FROM_EMAIL') or os.environ.get('SENDGRID_FROM_EMAIL')
    from_name  = os.environ.get('SMTP_FROM_NAME', 'JBJ Management')
    return from_email, from_name


def send_email_smtp(to_email, subject, html_body, attachments=None):
    """Send one email. Uses SendGrid HTTP API if SENDGRID_API_KEY is set
    (bypasses SMTP port blocking on cloud hosts); falls back to SMTP
    otherwise."""
    api_key = os.environ.get('SENDGRID_API_KEY')
    if api_key:
        _sendgrid_api_single(api_key, to_email, subject, html_body, attachments)
        return
    # SMTP fallback
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.application import MIMEApplication

    host = os.environ.get('SMTP_HOST')
    if not host:
        raise RuntimeError('Email sending is not configured on the server.')
    port      = int(os.environ.get('SMTP_PORT', '587'))
    username  = os.environ.get('SMTP_USERNAME')
    password  = os.environ.get('SMTP_PASSWORD')
    from_email, from_name = _sg_from()
    from_email = from_email or username
    use_tls = os.environ.get('SMTP_USE_TLS', 'true').strip().lower() not in ('0', 'false', 'no')

    msg = MIMEMultipart('mixed')
    msg['Subject'] = subject or '(no subject)'
    msg['From'] = f'{from_name} <{from_email}>'
    msg['To'] = to_email
    alt = MIMEMultipart('alternative')
    alt.attach(MIMEText(html_body, 'html'))
    msg.attach(alt)
    for filename, mimetype, data in (attachments or []):
        part = MIMEApplication(data, Name=filename)
        part['Content-Disposition'] = f'attachment; filename="{filename}"'
        msg.attach(part)
    with smtplib.SMTP(host, port, timeout=20) as server:
        if use_tls:
            server.starttls()
        if username and password:
            server.login(username, password)
        server.sendmail(from_email, [to_email], msg.as_string())


def _sendgrid_api_single(api_key, to_email, subject, html_body, attachments=None):
    """Send one email via SendGrid HTTP API (HTTPS, never blocked)."""
    import requests, base64
    from_email, from_name = _sg_from()
    if not from_email:
        raise RuntimeError('Set SMTP_FROM_EMAIL in environment variables.')

    payload = {
        'personalizations': [{'to': [{'email': to_email}]}],
        'from': {'email': from_email, 'name': from_name},
        'subject': subject or '(no subject)',
        'content': [{'type': 'text/html', 'value': html_body}],
    }
    if attachments:
        payload['attachments'] = [
            {'content': base64.b64encode(data).decode(), 'filename': fname,
             'type': mime, 'disposition': 'attachment'}
            for fname, mime, data in attachments
        ]
    res = requests.post(
        'https://api.sendgrid.com/v3/mail/send',
        headers={'Authorization': f'Bearer {api_key}'},
        json=payload, timeout=30,
    )
    if res.status_code not in (200, 202):
        raise RuntimeError(f'SendGrid error {res.status_code}: {res.text[:300]}')


def send_flyer_bulk_smtp(recipients, subject, html_body, png_bytes=None, png_filename='flyer.png'):
    """Bulk-send a flyer to a list of recipients. Uses SendGrid HTTP API
    when SENDGRID_API_KEY is set (BCC batches of 500); falls back to
    SMTP BCC batches of 50. Returns (sent_count, failed_count)."""
    api_key = os.environ.get('SENDGRID_API_KEY')
    if api_key:
        return _sendgrid_api_bulk(api_key, recipients, subject, html_body, png_bytes, png_filename)
    # SMTP fallback
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.application import MIMEApplication

    host = os.environ.get('SMTP_HOST')
    if not host:
        raise RuntimeError('Email sending is not configured on the server.')
    port      = int(os.environ.get('SMTP_PORT', '587'))
    username  = os.environ.get('SMTP_USERNAME')
    password  = os.environ.get('SMTP_PASSWORD')
    from_email, from_name = _sg_from()
    from_email = from_email or username
    use_tls = os.environ.get('SMTP_USE_TLS', 'true').strip().lower() not in ('0', 'false', 'no')

    BATCH = 50
    sent = failed = 0
    with smtplib.SMTP(host, port, timeout=30) as server:
        if use_tls:
            server.starttls()
        if username and password:
            server.login(username, password)
        for i in range(0, len(recipients), BATCH):
            batch = recipients[i:i + BATCH]
            msg = MIMEMultipart('mixed')
            msg['Subject'] = subject or '(no subject)'
            msg['From'] = f'{from_name} <{from_email}>'
            msg['To']   = f'{from_name} <{from_email}>'
            alt = MIMEMultipart('alternative')
            alt.attach(MIMEText(html_body, 'html'))
            msg.attach(alt)
            if png_bytes:
                part = MIMEApplication(png_bytes, Name=png_filename)
                part['Content-Disposition'] = f'attachment; filename="{png_filename}"'
                msg.attach(part)
            try:
                server.sendmail(from_email, [from_email] + batch, msg.as_string())
                sent += len(batch)
            except Exception:
                failed += len(batch)
    return sent, failed


def _sendgrid_api_bulk(api_key, recipients, subject, html_body, png_bytes=None, png_filename='flyer.png'):
    """Bulk-send via SendGrid HTTP API using BCC personalizations."""
    import requests, base64
    from_email, from_name = _sg_from()
    if not from_email:
        raise RuntimeError('Set SMTP_FROM_EMAIL in environment variables.')

    attachment = None
    if png_bytes:
        attachment = {'content': base64.b64encode(png_bytes).decode(),
                      'filename': png_filename, 'type': 'image/png',
                      'disposition': 'attachment'}

    BATCH = 500
    sent = failed = 0
    for i in range(0, len(recipients), BATCH):
        batch = recipients[i:i + BATCH]
        payload = {
            'personalizations': [{
                'to': [{'email': from_email}],
                'bcc': [{'email': e} for e in batch],
            }],
            'from': {'email': from_email, 'name': from_name},
            'subject': subject or '(no subject)',
            'content': [{'type': 'text/html', 'value': html_body}],
        }
        if attachment:
            payload['attachments'] = [attachment]
        try:
            res = requests.post(
                'https://api.sendgrid.com/v3/mail/send',
                headers={'Authorization': f'Bearer {api_key}'},
                json=payload, timeout=60,
            )
            if res.status_code in (200, 202):
                sent += len(batch)
            else:
                failed += len(batch)
        except Exception:
            failed += len(batch)
    return sent, failed


def _sendgrid_campaign_tracked(api_key, recipients, subject, html_body, send_id):
    """Send campaign with one personalization per recipient so SendGrid webhooks
    can include custom_args (send_id + contact_id) for per-contact tracking.
    Uses up to 1000 personalizations per request (SendGrid limit)."""
    import requests
    from_email, from_name = _sg_from()
    if not from_email:
        raise RuntimeError('Set SMTP_FROM_EMAIL in environment variables.')

    BATCH = 1000
    sent = failed = 0
    for i in range(0, len(recipients), BATCH):
        batch = recipients[i:i + BATCH]
        payload = {
            'personalizations': [
                {
                    'to': [{'email': r['email']}],
                    'custom_args': {
                        'send_id': str(send_id),
                        'contact_id': str(r['contact_id']),
                    },
                }
                for r in batch
            ],
            'from': {'email': from_email, 'name': from_name},
            'subject': subject or '(no subject)',
            'content': [{'type': 'text/html', 'value': html_body}],
            'tracking_settings': {
                'click_tracking': {'enable': True},
                'open_tracking': {'enable': True},
            },
        }
        try:
            res = requests.post(
                'https://api.sendgrid.com/v3/mail/send',
                headers={'Authorization': f'Bearer {api_key}'},
                json=payload, timeout=60,
            )
            if res.status_code in (200, 202):
                sent += len(batch)
            else:
                failed += len(batch)
        except Exception:
            failed += len(batch)
    return sent, failed


def _enroll_contact_in_sequences(contact_id, stage):
    """Enroll a contact in all active sequences triggered by the given stage,
    skipping if already enrolled and active."""
    from datetime import datetime as _dt, timedelta
    sequences = EmailSequence.query.filter_by(trigger_stage=stage, is_active=True).all()
    for seq in sequences:
        if not seq.steps:
            continue
        already = EmailSequenceEnrollment.query.filter_by(
            sequence_id=seq.id, contact_id=contact_id, status='active'
        ).first()
        if already:
            continue
        first_step = seq.steps[0]
        enrollment = EmailSequenceEnrollment(
            sequence_id=seq.id,
            contact_id=contact_id,
            next_step_index=0,
            next_send_at=_dt.utcnow() + timedelta(days=first_step.day_offset),
            status='active',
        )
        db.session.add(enrollment)
    db.session.commit()


def _process_sequence_emails(app):
    """Send any sequence emails that are due. Called by the scheduler."""
    import json as _json, urllib.request as _urlreq
    from datetime import datetime as _dt
    with app.app_context():
        now = _dt.utcnow()
        due = EmailSequenceEnrollment.query.filter(
            EmailSequenceEnrollment.status == 'active',
            EmailSequenceEnrollment.next_send_at <= now,
        ).all()
        sg_key = os.environ.get('SENDGRID_API_KEY') or (
            os.environ.get('SMTP_PASSWORD', '').startswith('SG.') and os.environ.get('SMTP_PASSWORD')
        ) or None
        from_email = os.environ.get('MAIL_FROM', os.environ.get('SMTP_USERNAME', ''))
        for enrollment in due:
            seq = enrollment.sequence
            if not seq.is_active:
                continue
            steps = seq.steps
            if enrollment.next_step_index >= len(steps):
                enrollment.status = 'completed'
                db.session.commit()
                continue
            step = steps[enrollment.next_step_index]
            contact = enrollment.contact
            if not contact or not contact.email:
                enrollment.status = 'cancelled'
                db.session.commit()
                continue
            # Send email
            if sg_key and from_email:
                try:
                    payload = {
                        'personalizations': [{'to': [{'email': contact.email,
                                                       'name': f'{contact.first_name or ""} {contact.last_name or ""}'.strip()}]}],
                        'from': {'email': from_email},
                        'subject': step.subject,
                        'content': [{'type': 'text/html', 'value': step.body}],
                    }
                    req = _urlreq.Request(
                        'https://api.sendgrid.com/v3/mail/send',
                        data=_json.dumps(payload).encode(),
                        headers={'Authorization': f'Bearer {sg_key}', 'Content-Type': 'application/json'},
                        method='POST',
                    )
                    _urlreq.urlopen(req, timeout=15)
                except Exception:
                    continue  # leave next_send_at as-is; retry next run
            # Advance to next step
            next_idx = enrollment.next_step_index + 1
            if next_idx >= len(steps):
                enrollment.status = 'completed'
                enrollment.next_step_index = next_idx
            else:
                next_step = steps[next_idx]
                from datetime import timedelta
                enrollment.next_step_index = next_idx
                enrollment.next_send_at = now + timedelta(days=next_step.day_offset)
            db.session.commit()


def create_app(config_class=Config):
    app = Flask(__name__)
    app.config.from_object(config_class)
    db.init_app(app)
    login_manager.init_app(app)
    limiter.init_app(app)

    with app.app_context():
        db.create_all()
        # Add columns that didn't exist in earlier schema versions
        for stmt in [
            "ALTER TABLE users ADD COLUMN IF NOT EXISTS can_post_social BOOLEAN NOT NULL DEFAULT FALSE",
            "ALTER TABLE tasks ADD COLUMN IF NOT EXISTS notes TEXT",
            "ALTER TABLE contacts ADD COLUMN IF NOT EXISTS pipeline_stage VARCHAR(32)",
            "ALTER TABLE flyer_templates ADD COLUMN background VARCHAR(32) DEFAULT '#ffffff'",
            "ALTER TABLE flyer_templates ADD COLUMN bg_asset_id INTEGER",
        ]:
            try:
                db.session.execute(db.text(stmt))
                db.session.commit()
            except Exception:
                db.session.rollback()
        _bootstrap_admin_user()

    # Start background scheduler for email sequences
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        _scheduler = BackgroundScheduler(daemon=True)
        _scheduler.add_job(
            _process_sequence_emails, 'interval', hours=1,
            args=[app], id='seq_emails', replace_existing=True,
        )
        _scheduler.start()
    except Exception:
        pass

    contact_schema = ContactSchema()
    contacts_schema = ContactSchema(many=True)

    @login_manager.user_loader
    def load_user(user_id):
        return User.query.get(int(user_id))

    # Every page and API route requires login except the login page itself
    # and static assets -- gated centrally here instead of decorating each
    # of the ~20 routes individually, so a newly added route can't
    # accidentally end up unprotected.
    @app.before_request
    def require_login():
        if request.endpoint in (None, 'login', 'static'):
            return None
        if not current_user.is_authenticated:
            if request.path.startswith('/api/'):
                return jsonify({'error': 'authentication required'}), 401
            return redirect(url_for('login', next=request.path))

    @app.route('/login', methods=['GET', 'POST'])
    @limiter.limit("10 per minute", methods=["POST"])
    def login():
        if current_user.is_authenticated:
            return redirect(url_for('index'))
        error = None
        username = ''
        if request.method == 'POST':
            username = (request.form.get('username') or '').strip()
            password = request.form.get('password') or ''
            user = User.query.filter_by(username=username).first()
            if user and user.check_password(password):
                login_user(user, remember=True)
                next_path = request.args.get('next')
                return redirect(next_path or url_for('index'))
            error = 'Invalid username or password.'
        return render_template('login.html', error=error, username=username)

    @app.errorhandler(429)
    def too_many_login_attempts(e):
        # Only /login is rate-limited today, so a generic handler is safe;
        # revisit this if another route ever gets its own limit.
        return render_template(
            'login.html',
            error='Too many login attempts. Please wait a minute and try again.',
            username=''
        ), 429

    @app.route('/logout')
    def logout():
        logout_user()
        return redirect(url_for('login'))

    @app.route('/')
    def index():
        # Serve single-page frontend
        return render_template('index.html')

    @app.route('/profile/<int:contact_id>')
    def profile(contact_id):
        return render_template('profile.html', contact_id=contact_id)

    # ── Proposal Manager ──────────────────────────────────────────────── #

    @app.route('/proposals')
    @login_required
    def proposals_hub():
        proposals = Proposal.query.order_by(Proposal.updated_at.desc()).all()
        cs_count = CaseStudy.query.count()
        contact_count = Contact.query.count()
        return render_template('proposals_hub.html',
                               proposals=proposals,
                               cs_count=cs_count,
                               contact_count=contact_count)

    @app.route('/proposals/new')
    @login_required
    def proposal_new():
        case_studies = CaseStudy.query.order_by(CaseStudy.created_at.desc()).all()
        return render_template('proposal_builder.html',
                               proposal=None,
                               case_studies=case_studies)

    @app.route('/proposals/<int:proposal_id>')
    @login_required
    def proposal_detail(proposal_id):
        proposal = Proposal.query.get_or_404(proposal_id)
        case_studies = CaseStudy.query.order_by(CaseStudy.created_at.desc()).all()
        # Hydrate linked contacts
        contact_ids = proposal.contact_ids or []
        contacts = Contact.query.filter(Contact.id.in_(contact_ids)).all() if contact_ids else []
        # Hydrate linked case studies
        cs_ids = proposal.case_study_ids or []
        linked_cs = CaseStudy.query.filter(CaseStudy.id.in_(cs_ids)).all() if cs_ids else []
        return render_template('proposal_builder.html',
                               proposal=proposal,
                               case_studies=case_studies,
                               linked_contacts=contacts,
                               linked_cs=linked_cs)

    @app.route('/proposals/list')
    @login_required
    def proposals_list():
        status_filter = request.args.get('status', '')
        q = Proposal.query.order_by(Proposal.updated_at.desc())
        if status_filter:
            q = q.filter(Proposal.status == status_filter)
        proposals = q.all()
        return render_template('proposals_list.html',
                               proposals=proposals,
                               status_filter=status_filter)

    @app.route('/api/proposals', methods=['GET'])
    @login_required
    def list_proposals():
        per_page = request.args.get('per_page', 50, type=int)
        page = request.args.get('page', 1, type=int)
        q = Proposal.query.order_by(Proposal.updated_at.desc())
        total = q.count()
        items = q.offset((page - 1) * per_page).limit(per_page).all()
        return jsonify({'proposals': [p.to_dict() for p in items], 'total': total, 'per_page': per_page})

    @app.route('/api/case-studies-count')
    @login_required
    def case_studies_count():
        return jsonify({'count': CaseStudy.query.count()})

    @app.route('/api/proposals', methods=['POST'])
    @login_required
    def create_proposal():
        data = request.get_json(force=True)
        p = Proposal(
            title=data.get('title', 'Untitled Proposal'),
            client_name=data.get('client_name'),
            client_org=data.get('client_org'),
            contact_ids=data.get('contact_ids', []),
            case_study_ids=data.get('case_study_ids', []),
            overview=data.get('overview'),
            scope=data.get('scope'),
            timeline=data.get('timeline'),
            budget=data.get('budget'),
            notes=data.get('notes'),
            status=data.get('status', 'draft'),
            created_by_id=current_user.id,
        )
        db.session.add(p)
        db.session.commit()
        return jsonify(p.to_dict()), 201

    @app.route('/api/proposals/<int:proposal_id>', methods=['PUT'])
    @login_required
    def update_proposal(proposal_id):
        p = Proposal.query.get_or_404(proposal_id)
        data = request.get_json(force=True)
        for field in ('title', 'client_name', 'client_org', 'contact_ids',
                      'case_study_ids', 'overview', 'scope', 'timeline',
                      'budget', 'notes', 'status'):
            if field in data:
                setattr(p, field, data[field])
        db.session.commit()
        return jsonify(p.to_dict())

    @app.route('/api/proposals/<int:proposal_id>', methods=['DELETE'])
    @login_required
    def delete_proposal(proposal_id):
        p = Proposal.query.get_or_404(proposal_id)
        db.session.delete(p)
        db.session.commit()
        return jsonify({'ok': True})

    @app.route('/api/proposals/generate', methods=['POST'])
    @login_required
    def generate_proposal():
        import anthropic
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            return jsonify({'error': 'ANTHROPIC_API_KEY is not configured on the server.'}), 500

        data = request.get_json(force=True)
        title       = (data.get('title') or '').strip()
        client_name = (data.get('client_name') or '').strip()
        client_org  = (data.get('client_org') or '').strip()
        timeline    = (data.get('timeline') or '').strip()
        budget      = (data.get('budget') or '').strip()
        user_prompt = (data.get('prompt') or '').strip()
        cs_ids      = data.get('case_study_ids') or []

        # Build context block
        lines = ['You are drafting a project proposal for JBJ Management.']
        if title:        lines.append(f'Proposal title: {title}')
        if client_name:  lines.append(f'Client: {client_name}')
        if client_org:   lines.append(f'Organization: {client_org}')
        if timeline:     lines.append(f'Timeline: {timeline}')
        if budget:       lines.append(f'Budget: {budget}')

        # Pull case study details
        if cs_ids:
            case_studies = CaseStudy.query.filter(CaseStudy.id.in_(cs_ids)).all()
            for cs in case_studies:
                block = [f'\nRelevant past work - "{cs.title}"']
                if cs.client:  block.append(f'Client: {cs.client}')
                if cs.sector:  block.append(f'Sector: {cs.sector}')
                text = '\n'.join(filter(None, [cs.challenges, cs.solution, cs.results])) or cs.extracted_text or ''
                if text: block.append(text[:1500])
                lines.extend(block)

        context = '\n'.join(lines)
        user_content = context
        if user_prompt:
            user_content += f'\n\nAdditional context from the user: {user_prompt}'

        system = (
            'You write professional project proposals for JBJ Management, a talent and project management company. '
            'Given the context below, write two sections:\n'
            '1. A concise "Overview" paragraph (3-5 sentences) that summarizes the project and its value to the client.\n'
            '2. A "Scope of Work" section (4-8 bullet points) that details deliverables, services, and milestones.\n'
            'Be specific and professional. Do not invent facts not provided - use [PLACEHOLDER] for missing specifics. '
            'Respond ONLY with a JSON object: {"overview": "...", "scope": "..."}'
        )

        try:
            client = anthropic.Anthropic(api_key=api_key)
            response = client.messages.create(
                model='claude-sonnet-4-6',
                max_tokens=1200,
                system=system,
                messages=[{'role': 'user', 'content': user_content}],
            )
            text = next((b.text for b in response.content if b.type == 'text'), '{}')
            import json as _json
            # Strip markdown fences if present
            text = text.strip()
            if text.startswith('```'):
                text = text.split('```')[1]
                if text.startswith('json'): text = text[4:]
            result = _json.loads(text.strip())
            return jsonify({'overview': result.get('overview', ''), 'scope': result.get('scope', '')})
        except anthropic.APIStatusError as e:
            return jsonify({'error': f'Claude API error: {e.message}'}), 502
        except Exception as e:
            return jsonify({'error': str(e)}), 502

    @app.route('/case-studies')
    def case_studies_list():
        q = request.args.get('q', type=str)
        sector = request.args.get('sector', type=str)
        page = request.args.get('page', default=1, type=int)
        limit = 12

        query = CaseStudy.query
        if q:
            like = f"%{q}%"
            query = query.filter(or_(
                CaseStudy.title.ilike(like),
                CaseStudy.client.ilike(like),
                CaseStudy.sector.ilike(like),
                CaseStudy.challenges.ilike(like),
                CaseStudy.solution.ilike(like),
                CaseStudy.results.ilike(like),
                CaseStudy.extracted_text.ilike(like),
            ))
        if sector:
            query = query.filter(CaseStudy.sector == sector)

        total = query.count()
        pages = max(1, (total + limit - 1) // limit)
        page = max(1, min(page, pages))
        items = query.order_by(CaseStudy.created_at.desc()).offset((page - 1) * limit).limit(limit).all()
        sectors = [s[0] for s in db.session.query(CaseStudy.sector)
                   .filter(CaseStudy.sector.isnot(None), CaseStudy.sector != '')
                   .distinct().order_by(CaseStudy.sector).all()]

        return render_template(
            'case_studies.html', items=items, total=total, page=page, pages=pages,
            q=q or '', sector=sector or '', sectors=sectors,
        )

    @app.route('/case-studies/new')
    def case_study_new_form():
        if not current_user.is_admin:
            return redirect(url_for('case_studies_list'))
        return render_template('case_study_form.html', case_study=None)

    @app.route('/case-studies/upload')
    def case_study_upload_form():
        if not current_user.is_admin:
            return redirect(url_for('case_studies_list'))
        return render_template('case_study_upload.html')

    @app.route('/case-studies/<int:case_study_id>')
    def case_study_detail(case_study_id):
        cs = CaseStudy.query.get_or_404(case_study_id)
        return render_template('case_study_detail.html', cs=cs)

    @app.route('/case-studies/<int:case_study_id>/file')
    def case_study_file(case_study_id):
        cs = CaseStudy.query.get_or_404(case_study_id)
        if not cs.file_data:
            abort(404)
        return send_file(
            io.BytesIO(cs.file_data),
            mimetype=cs.file_mimetype or 'application/octet-stream',
            as_attachment=True,
            download_name=cs.file_name or f'case-study-{cs.id}',
        )

    @app.route('/case-studies/<int:case_study_id>/edit')
    def case_study_edit_form(case_study_id):
        if not current_user.is_admin:
            return redirect(url_for('case_studies_list'))
        cs = CaseStudy.query.get_or_404(case_study_id)
        return render_template('case_study_form.html', case_study=cs)

    @app.route('/admin')
    def admin():
        if not current_user.is_admin:
            return redirect(url_for('index'))
        return render_template('admin.html')

    @app.route('/admin/sync')
    def admin_sync():
        if not current_user.is_admin:
            return redirect(url_for('index'))
        return render_template('admin_sync.html')

    @app.route('/admin/users')
    def manage_users():
        if not current_user.is_admin:
            return redirect(url_for('index'))
        users = User.query.order_by(User.username).all()
        social = {t.platform: t for t in SocialToken.query.all()}
        li_configured = bool(_social_cfg('LINKEDIN_CLIENT_ID'))
        fb_configured = bool(_social_cfg('FACEBOOK_APP_ID'))
        return render_template('users.html', users=users, social=social,
                               li_configured=li_configured, fb_configured=fb_configured)

    @app.route('/admin/audit')
    def audit_log():
        if not current_user.is_admin:
            return redirect(url_for('index'))
        page = request.args.get('page', default=1, type=int)
        limit = 50
        query = AuditLog.query.order_by(AuditLog.created_at.desc())
        total = query.count()
        rows = query.offset((page - 1) * limit).limit(limit).all()
        pages = max(1, (total + limit - 1) // limit)
        entries = [{
            'created_at': e.created_at,
            'actor_name': e.actor_name,
            'summary': format_audit_summary(e),
        } for e in rows]
        return render_template('audit.html', entries=entries, page=page, pages=pages)

    @app.route('/admin/analytics')
    def analytics():
        if not current_user.is_admin:
            return redirect(url_for('index'))

        def ranked(rows):
            """[(label, count), ...] -> [{'label','count','pct'}, ...], pct relative to the top row."""
            top = rows[0][1] if rows else 0
            return [{'label': label, 'count': count, 'pct': round(100 * count / top) if top else 0}
                    for label, count in rows]

        total_contacts = Contact.query.count()
        total_orgs = OutreachOrg.query.count()
        total_activities = Activity.query.count()
        incomplete_count = Contact.query.filter(contact_incomplete_clause()).count()
        data_complete_pct = round(100 * (total_contacts - incomplete_count) / total_contacts) if total_contacts else 0

        contacted_ids = {r[0] for r in db.session.query(Activity.contact_id)
                          .filter(Activity.contact_id.isnot(None)).distinct().all()}
        never_contacted = max(0, total_contacts - len(contacted_ids))

        today = date.today()
        activities_30d = Activity.query.filter(Activity.contacted_on >= today - timedelta(days=30)).count()
        audit_cutoff = datetime.utcnow() - timedelta(days=30)
        audit_30d = AuditLog.query.filter(AuditLog.created_at >= audit_cutoff).count()

        # Weekly outreach trend, oldest to newest, last 12 weeks.
        weeks_back = 12
        trend_start = today - timedelta(weeks=weeks_back)
        contacted_dates = [d for (d,) in db.session.query(Activity.contacted_on)
                           .filter(Activity.contacted_on >= trend_start).all() if d]
        week_counts = [0] * weeks_back
        for d in contacted_dates:
            idx = weeks_back - 1 - (today - d).days // 7
            if 0 <= idx < weeks_back:
                week_counts[idx] += 1
        max_week = max(week_counts) if week_counts else 0
        weekly_trend = []
        for idx, c in enumerate(week_counts):
            weeks_ago = weeks_back - 1 - idx
            week_end = today - timedelta(days=weeks_ago * 7)
            week_start = week_end - timedelta(days=6)
            weekly_trend.append({
                'count': c,
                'pct': round(100 * c / max_week) if max_week else 0,
                'week_start': week_start.isoformat(),
                'week_end': week_end.isoformat(),
            })

        by_employee = ranked(
            db.session.query(Activity.employee_name, func.count(Activity.id))
            .group_by(Activity.employee_name)
            .order_by(func.count(Activity.id).desc()).limit(10).all()
        )

        by_channel = ranked(
            db.session.query(Activity.channel, func.count(Activity.id))
            .group_by(Activity.channel)
            .order_by(func.count(Activity.id).desc()).all()
        )
        for row in by_channel:
            row['label'] = row['label'] or 'Unspecified'

        # Contact.county can hold several comma-separated counties in one
        # field (see /api/counties) -- split those apart so each county
        # name is counted on its own instead of every combination being a
        # separate bucket.
        county_counts = {}
        county_rows = (db.session.query(Contact.county, func.count(Activity.id))
                        .join(Activity, Activity.contact_id == Contact.id)
                        .filter(Contact.county.isnot(None), Contact.county != '')
                        .group_by(Contact.county).all())
        for county_str, count in county_rows:
            for part in county_str.split(','):
                part = part.strip()
                if part:
                    county_counts[part] = county_counts.get(part, 0) + count
        by_county = ranked(sorted(county_counts.items(), key=lambda kv: -kv[1])[:10])

        by_action = ranked(
            db.session.query(AuditLog.action, func.count(AuditLog.id))
            .filter(AuditLog.created_at >= audit_cutoff)
            .group_by(AuditLog.action)
            .order_by(func.count(AuditLog.id).desc()).all()
        )
        for row in by_action:
            row['action'] = row['label']
            row['label'] = ACTION_LABELS.get(row['label'], row['label'])

        return render_template(
            'analytics.html',
            total_contacts=total_contacts,
            total_orgs=total_orgs,
            total_activities=total_activities,
            activities_30d=activities_30d,
            never_contacted=never_contacted,
            data_complete_pct=data_complete_pct,
            audit_30d=audit_30d,
            weekly_trend=weekly_trend,
            by_employee=by_employee,
            by_channel=by_channel,
            by_county=by_county,
            by_action=by_action,
        )

    @app.route('/api/analytics/activities', methods=['GET'])
    def analytics_activities():
        """Backs the click-to-drill-down on the Analytics dashboard --
        given one of the breakdown dimensions (employee/channel/county) or a
        week range from the trend chart, returns the actual outreach entries
        behind that number, including who the contact was."""
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403

        employee = request.args.get('employee', type=str)
        channel = request.args.get('channel', type=str)
        county = request.args.get('county', type=str)
        week_start = request.args.get('week_start', type=str)
        week_end = request.args.get('week_end', type=str)

        query = Activity.query.options(db.joinedload(Activity.contact))
        if employee:
            query = query.filter(Activity.employee_name == employee)
        if channel:
            if channel == 'Unspecified':
                query = query.filter(or_(Activity.channel.is_(None), Activity.channel == ''))
            else:
                query = query.filter(Activity.channel == channel)
        if county:
            query = query.join(Contact, Activity.contact_id == Contact.id).filter(Contact.county.ilike(f'%{county}%'))
        if week_start and week_end:
            query = query.filter(Activity.contacted_on >= week_start, Activity.contacted_on <= week_end)

        rows = query.order_by(Activity.contacted_on.desc()).limit(200).all()
        activities = []
        for a in rows:
            contact = a.contact
            contact_name = f"{contact.first_name or ''} {contact.last_name or ''}".strip() if contact else ''
            activities.append({
                'employee_name': a.employee_name,
                'channel': a.channel,
                'contacted_on': a.contacted_on.isoformat() if a.contacted_on else None,
                'summary': a.summary,
                'contact_name': contact_name or None,
                'contact_email': contact.email if contact else None,
                'organization': a.organization or (contact.organization if contact else None),
            })
        return jsonify({'count': len(activities), 'activities': activities})

    @app.route('/api/analytics/audit-entries', methods=['GET'])
    def analytics_audit_entries():
        """Backs the click-to-drill-down on the Analytics dashboard's Admin
        Activity breakdown -- same idea as analytics_activities(), but for
        AuditLog rows instead of outreach Activity rows."""
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403

        action = request.args.get('action', type=str)
        query = AuditLog.query.filter(AuditLog.created_at >= datetime.utcnow() - timedelta(days=30))
        if action:
            query = query.filter(AuditLog.action == action)

        rows = query.order_by(AuditLog.created_at.desc()).limit(200).all()
        entries = [{
            'created_at': e.created_at.isoformat() if e.created_at else None,
            'actor_name': e.actor_name,
            'summary': format_audit_summary(e),
        } for e in rows]
        return jsonify({'count': len(entries), 'entries': entries})

    @app.route('/api/users', methods=['POST'])
    def create_user_api():
        # Lets an admin create more employee logins from the browser, so
        # only the very first account ever needs the env-var bootstrap above.
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403
        data = request.get_json(force=True) or {}
        username = (data.get('username') or '').strip()
        password = data.get('password') or ''
        display_name = (data.get('display_name') or '').strip() or username
        is_admin = bool(data.get('is_admin'))
        if not username or not password:
            return jsonify({'error': 'username and password are required'}), 400
        if len(password) < 8:
            return jsonify({'error': 'password must be at least 8 characters'}), 400
        if User.query.filter_by(username=username).first():
            return jsonify({'error': 'that username is already taken'}), 409
        user = User(username=username, display_name=display_name, is_admin=is_admin)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        log_audit('user_created', 'user', user.id, username, {'is_admin': is_admin})
        return jsonify(user.to_dict()), 201

    @app.route('/api/users/<int:user_id>/reset-password', methods=['POST'])
    def reset_user_password(user_id):
        # There's no self-service "forgot password" flow (no email service
        # configured) -- an admin resetting it here is the only path today.
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403
        data = request.get_json(silent=True) or {}
        password = data.get('password') or ''
        if len(password) < 8:
            return jsonify({'error': 'password must be at least 8 characters'}), 400
        user = User.query.get_or_404(user_id)
        user.set_password(password)
        db.session.commit()
        log_audit('password_reset', 'user', user.id, user.username)
        return jsonify({'ok': True})

    @app.route('/api/users/<int:user_id>/toggle-social', methods=['POST'])
    @login_required
    def toggle_user_social(user_id):
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403
        user = User.query.get_or_404(user_id)
        user.can_post_social = not user.can_post_social
        db.session.commit()
        log_audit('social_permission_changed', 'user', user.id, user.username,
                  {'can_post_social': user.can_post_social})
        return jsonify({'ok': True, 'can_post_social': user.can_post_social})

    def _case_study_fields(data):
        return {
            'title': (data.get('title') or '').strip(),
            'client': (data.get('client') or '').strip() or None,
            'sector': (data.get('sector') or '').strip() or None,
            'challenges': (data.get('challenges') or '').strip() or None,
            'solution': (data.get('solution') or '').strip() or None,
            'results': (data.get('results') or '').strip() or None,
        }

    @app.route('/api/case-studies', methods=['GET'])
    def list_case_studies_json():
        """Lightweight list (no full text) for pickers like Draft Email's
        case-study reference dropdown -- the HTML /case-studies page has its
        own richer, paginated query and isn't reused here."""
        items = CaseStudy.query.order_by(CaseStudy.title).all()
        return jsonify({'case_studies': [
            {'id': c.id, 'title': c.title, 'client': c.client, 'sector': c.sector}
            for c in items
        ]})

    @app.route('/api/case-studies', methods=['POST'])
    def create_case_study():
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403
        fields = _case_study_fields(request.get_json(force=True) or {})
        if not fields['title']:
            return jsonify({'error': 'title is required'}), 400
        cs = CaseStudy(**fields)
        db.session.add(cs)
        db.session.commit()
        log_audit('case_study_created', 'case_study', cs.id, cs.title)
        return jsonify(cs.to_dict()), 201

    @app.route('/api/case-studies/<int:case_study_id>', methods=['PUT'])
    def update_case_study(case_study_id):
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403
        cs = CaseStudy.query.get_or_404(case_study_id)
        fields = _case_study_fields(request.get_json(force=True) or {})
        if not fields['title']:
            return jsonify({'error': 'title is required'}), 400
        for key, value in fields.items():
            setattr(cs, key, value)
        db.session.commit()
        log_audit('case_study_updated', 'case_study', cs.id, cs.title)
        return jsonify(cs.to_dict())

    @app.route('/api/case-studies/<int:case_study_id>', methods=['DELETE'])
    def delete_case_study(case_study_id):
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403
        cs = CaseStudy.query.get_or_404(case_study_id)
        label = cs.title
        db.session.delete(cs)
        db.session.commit()
        log_audit('case_study_deleted', 'case_study', case_study_id, label)
        return jsonify({'deleted': True})

    @app.route('/api/case-studies/upload', methods=['POST'])
    def upload_case_study_files():
        """Stores uploaded PDF/Word files as-is (downloadable later) and
        best-effort extracts their text for search -- no AI involved.
        Title defaults to the filename; sector/client/challenges/solution/
        results are left blank for the admin to fill in later via Edit if
        they want that level of detail."""
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403

        files = request.files.getlist('files')
        if not files:
            return jsonify({'error': 'No files uploaded.'}), 400
        if len(files) > 15:
            return jsonify({'error': 'Upload at most 15 files at a time.'}), 400

        results = []
        created_titles = []
        for f in files:
            filename = f.filename or 'Untitled'
            file_bytes = f.read()
            if not file_bytes:
                results.append({'filename': filename, 'success': False, 'error': 'Empty file.'})
                continue
            f.stream.seek(0)
            text, extract_error = extract_case_study_text(f)

            title = filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').strip() or filename
            cs = CaseStudy(
                title=title,
                file_data=file_bytes,
                file_name=filename,
                file_mimetype=f.mimetype or 'application/octet-stream',
                extracted_text=text,
            )
            db.session.add(cs)
            db.session.commit()
            created_titles.append(cs.title)
            results.append({
                'filename': filename, 'success': True, 'id': cs.id, 'title': cs.title,
                'text_extracted': bool(text),
                'note': None if text else (extract_error or 'File stored, but text could not be extracted for search.'),
            })

        if created_titles:
            log_audit('case_study_uploaded', 'case_study', None, None, {'count': len(created_titles), 'titles': created_titles})

        return jsonify({'results': results})

    @app.route('/email-events')
    @login_required
    def email_events_hub():
        return render_template('email_events_hub.html')

    @app.route('/email-builder')
    def email_builder_list():
        """Saved email designs from the drag-and-drop builder. Open to
        everyone logged in, same as Draft Email/Create Flyer -- this is a
        drafting tool, not a sensitive admin page. Sending (a later step)
        is where a real-people-get-emailed confirmation belongs, not here."""
        templates = EmailTemplate.query.filter(
            db.or_(EmailTemplate.is_public == True, EmailTemplate.created_by_id == current_user.id)
        ).order_by(EmailTemplate.updated_at.desc()).all()
        return render_template('email_builder_list.html', templates=templates, current_user_id=current_user.id)

    @app.route('/email-builder/<int:template_id>')
    def email_builder_edit(template_id):
        template = EmailTemplate.query.get_or_404(template_id)
        return render_template('email_builder.html', template=template)

    @app.route('/api/email-templates', methods=['GET'])
    def list_email_templates():
        items = EmailTemplate.query.filter(
            db.or_(EmailTemplate.is_public == True, EmailTemplate.created_by_id == current_user.id)
        ).order_by(EmailTemplate.updated_at.desc()).all()
        return jsonify({'email_templates': [t.to_dict() for t in items]})

    @app.route('/api/email-templates', methods=['POST'])
    def create_email_template():
        data = request.get_json(force=True) or {}
        name = (data.get('name') or '').strip() or 'Untitled email'
        t = EmailTemplate(
            name=name,
            subject=(data.get('subject') or '').strip() or None,
            blocks=data.get('blocks') or [],
            created_by_id=current_user.id,
        )
        db.session.add(t)
        db.session.commit()
        log_audit('email_template_created', 'email_template', t.id, t.name)
        return jsonify(t.to_dict()), 201

    @app.route('/api/email-templates/<int:template_id>', methods=['GET'])
    def get_email_template(template_id):
        t = EmailTemplate.query.get_or_404(template_id)
        return jsonify(t.to_dict())

    @app.route('/api/email-templates/<int:template_id>', methods=['PUT'])
    def update_email_template(template_id):
        t = EmailTemplate.query.get_or_404(template_id)
        data = request.get_json(force=True) or {}
        t.name = (data.get('name') or '').strip() or 'Untitled email'
        t.subject = (data.get('subject') or '').strip() or None
        t.blocks = data.get('blocks') or []
        if 'is_public' in data:
            t.is_public = bool(data['is_public'])
        db.session.commit()
        log_audit('email_template_updated', 'email_template', t.id, t.name)
        return jsonify(t.to_dict())

    @app.route('/api/email-templates/<int:template_id>', methods=['DELETE'])
    def delete_email_template(template_id):
        t = EmailTemplate.query.get_or_404(template_id)
        label = t.name
        db.session.delete(t)
        db.session.commit()
        log_audit('email_template_deleted', 'email_template', template_id, label)
        return jsonify({'deleted': True})

    @app.route('/api/email-templates/<int:template_id>/duplicate', methods=['POST'])
    def duplicate_email_template(template_id):
        import copy as _copy
        t = EmailTemplate.query.get_or_404(template_id)
        dup = EmailTemplate(
            name=f'{t.name} (copy)',
            subject=t.subject,
            blocks=_copy.deepcopy(t.blocks),
            is_public=False,
            created_by_id=current_user.id,
        )
        db.session.add(dup)
        db.session.commit()
        log_audit('email_template_created', 'email_template', dup.id, dup.name)
        return jsonify(dup.to_dict()), 201

    @app.route('/api/email-templates/<int:template_id>/send-bulk', methods=['POST'])
    def send_email_template_bulk(template_id):
        """Send the email template to a filtered group of contacts.
        Merge tags ({{first_name}} etc.) are replaced per-contact.
        Only contacts with a non-empty email address receive the email."""
        import re as _re
        t = EmailTemplate.query.get_or_404(template_id)
        data = request.get_json(force=True) or {}
        tag_filter = (data.get('tag') or '').strip()
        subject_override = (data.get('subject') or '').strip()
        html_body = (data.get('html') or '').strip()
        if not html_body:
            return jsonify({'error': 'Email has no content.'}), 400

        query = Contact.query.filter(Contact.email.isnot(None), Contact.email != '')
        if tag_filter:
            query = query.filter(Contact.tag == tag_filter)
        contacts = query.all()
        if not contacts:
            return jsonify({'error': 'No contacts with email addresses match that filter.'}), 400

        subject = subject_override or t.subject or t.name or 'Email from JBJ Management'

        def replace_tags(text, contact):
            replacements = {
                '{{first_name}}': contact.first_name or '',
                '{{last_name}}':  contact.last_name or '',
                '{{full_name}}':  ' '.join(filter(None, [contact.first_name, contact.last_name])),
                '{{organization}}': contact.organization or '',
                '{{title}}':      contact.title or '',
            }
            for tag, val in replacements.items():
                text = text.replace(tag, val)
            return text

        host_url = request.host_url
        sent, failed = 0, 0
        errors = []
        template_id_value, template_name = t.id, t.name
        db.session.close()

        for contact in contacts:
            try:
                body = replace_tags(html_body, contact)
                body = absolutize_static_urls(body, host_url)
                wrapped = (
                    '<div style="font-family:Arial,Helvetica,sans-serif;max-width:600px;'
                    f'margin:0 auto;padding:16px;">{body}</div>'
                )
                contact_subject = replace_tags(subject, contact)
                send_email_smtp(contact.email, contact_subject, wrapped, [])
                sent += 1
            except Exception as e:
                failed += 1
                errors.append(str(e))

        log_audit('email_template_sent', 'email_template', template_id_value, template_name, {
            'bulk': True, 'tag': tag_filter, 'sent': sent, 'failed': failed,
        })
        return jsonify({'sent': sent, 'failed': failed, 'errors': errors[:5]})

    @app.route('/api/email-templates/<int:template_id>/send', methods=['POST'])
    def send_email_template(template_id):
        """Sends the current compose content to one typed-in address --
        a quick send/test capability, not the bulk "send to my filtered
        contact list" feature (that needs batching, an unsubscribe
        mechanism, and background dispatch since it could be hundreds of
        real people; this is a single address, fast enough to send inline
        within the request)."""
        t = EmailTemplate.query.get_or_404(template_id)
        to_email = (request.form.get('to') or '').strip()
        if not to_email:
            return jsonify({'error': 'Enter a recipient email address.'}), 400
        subject = (request.form.get('subject') or t.subject or t.name or '').strip()
        html_body = request.form.get('html') or ''
        if not html_body.strip():
            return jsonify({'error': 'This email has no content yet.'}), 400

        files = request.files.getlist('attachments')
        if len(files) > 5:
            return jsonify({'error': 'Attach at most 5 files.'}), 400
        attachments = []
        total_bytes = 0
        for f in files:
            if not f or not f.filename:
                continue
            data = f.read()
            total_bytes += len(data)
            if total_bytes > 15 * 1024 * 1024:
                return jsonify({'error': 'Attachments are too large (15MB total limit).'}), 400
            attachments.append((f.filename, f.mimetype or 'application/octet-stream', data))

        html_body = absolutize_static_urls(html_body, request.host_url)
        wrapped_html = (
            '<div style="font-family:Arial,Helvetica,sans-serif;max-width:600px;'
            f'margin:0 auto;padding:16px;">{html_body}</div>'
        )

        # Capture what log_audit needs as plain values, then release the DB
        # connection before the slow SMTP network call. Flask-SQLAlchemy
        # holds one connection checked out for the whole request -- left
        # idle for the several seconds an SMTP handshake can take, Neon's
        # pooler closes it server-side, and the next query after (the
        # audit log write, or Flask-Login re-touching current_user) fails
        # with "SSL connection has been closed unexpectedly" even though
        # the email itself sent fine. db.session.close() forces a fresh
        # connection to be checked out afterward instead of reusing the
        # one that just sat idle through the delay.
        template_id_value, template_name = t.id, t.name
        db.session.close()

        try:
            send_email_smtp(to_email, subject, wrapped_html, attachments)
        except RuntimeError as e:
            return jsonify({'error': str(e)}), 500
        except Exception as e:
            return jsonify({'error': f'Could not send: {e}'}), 502

        log_audit('email_template_sent', 'email_template', template_id_value, template_name, {
            'to': to_email, 'attachment_count': len(attachments),
        })
        return jsonify({'sent': True})

    # ------------------------------------------------------------------ #
    # Flyer / canvas builder                                               #
    # ------------------------------------------------------------------ #

    def _valid_flyer_formats():
        from flyer_render import CANVAS_FORMATS
        return set(CANVAS_FORMATS.keys())

    @app.route('/flyer-builder')
    def flyer_builder_list():
        templates = FlyerTemplate.query.filter(
            db.or_(FlyerTemplate.is_public == True, FlyerTemplate.created_by_id == current_user.id)
        ).order_by(FlyerTemplate.updated_at.desc()).all()
        return render_template('flyer_builder_list.html', templates=templates, current_user_id=current_user.id)

    @app.route('/flyer-builder/<int:template_id>')
    def flyer_builder_edit(template_id):
        template = FlyerTemplate.query.get_or_404(template_id)
        return render_template('flyer_builder.html', template=template)

    @app.route('/api/flyer-templates', methods=['GET'])
    def list_flyer_templates():
        items = FlyerTemplate.query.filter(
            db.or_(FlyerTemplate.is_public == True, FlyerTemplate.created_by_id == current_user.id)
        ).order_by(FlyerTemplate.updated_at.desc()).all()
        return jsonify({'flyer_templates': [t.to_dict() for t in items]})

    @app.route('/api/flyer-templates', methods=['POST'])
    def create_flyer_template():
        data = request.get_json(force=True) or {}
        t = FlyerTemplate(
            name=(data.get('name') or '').strip() or 'Untitled flyer',
            format=data.get('format', 'square') if data.get('format') in _valid_flyer_formats() else 'square',
            elements=data.get('elements') or [],
            background=data.get('background') or '#ffffff',
            bg_asset_id=data.get('bg_asset_id'),
            created_by_id=current_user.id,
        )
        db.session.add(t)
        db.session.commit()
        log_audit('flyer_template_created', 'flyer_template', t.id, t.name)
        return jsonify(t.to_dict()), 201

    @app.route('/api/flyer-templates/<int:template_id>', methods=['GET'])
    def get_flyer_template(template_id):
        return jsonify(FlyerTemplate.query.get_or_404(template_id).to_dict())

    @app.route('/api/flyer-templates/<int:template_id>', methods=['PUT'])
    def update_flyer_template(template_id):
        t = FlyerTemplate.query.get_or_404(template_id)
        data = request.get_json(force=True) or {}
        t.name = (data.get('name') or '').strip() or 'Untitled flyer'
        if data.get('format') in _valid_flyer_formats():
            t.format = data['format']
        t.elements = data.get('elements') or []
        if 'background' in data:
            t.background = (data.get('background') or '#ffffff')[:32]
        if 'bg_asset_id' in data:
            t.bg_asset_id = data.get('bg_asset_id')
        if 'is_public' in data:
            t.is_public = bool(data['is_public'])
        db.session.commit()
        log_audit('flyer_template_updated', 'flyer_template', t.id, t.name)
        return jsonify(t.to_dict())

    @app.route('/api/flyer-templates/<int:template_id>', methods=['DELETE'])
    def delete_flyer_template(template_id):
        t = FlyerTemplate.query.get_or_404(template_id)
        label = t.name
        db.session.delete(t)
        db.session.commit()
        log_audit('flyer_template_deleted', 'flyer_template', template_id, label)
        return jsonify({'deleted': True})

    @app.route('/api/flyer-templates/<int:template_id>/send', methods=['POST'])
    @login_required
    def send_flyer_template(template_id):
        t = FlyerTemplate.query.get_or_404(template_id)
        data = request.get_json(force=True) or {}

        subject         = (data.get('subject') or '').strip() or t.name
        message         = (data.get('message') or '').strip()
        tag             = data.get('tag', '')
        county          = data.get('county', '')
        q               = data.get('q', '')
        followup        = data.get('followup', '')
        favorites_only  = bool(data.get('favorites_only', False))
        background      = data.get('background', t.background or '#ffffff')
        bg_asset_id     = data.get('bg_asset_id', t.bg_asset_id)

        contacts = filtered_contacts_query(
            q=q, tag=tag, county=county, followup=followup, favorites_only=favorites_only
        ).filter(
            Contact.email.isnot(None),
            Contact.email != '',
            Contact.unsubscribed == False,
        ).all()

        if not contacts:
            return jsonify({'error': 'No contacts with email addresses match this filter.'}), 400

        from flyer_render import render_flyer_png

        def asset_loader(asset_id):
            try:
                a = FlyerAsset.query.get(int(asset_id))
                return a.data if a else None
            except (TypeError, ValueError):
                return None

        bg_img_bytes = None
        if bg_asset_id:
            try:
                _a = FlyerAsset.query.get(int(bg_asset_id))
                if _a:
                    bg_img_bytes = _a.data
            except (TypeError, ValueError):
                pass
        png_bytes = render_flyer_png(t.elements, fmt=t.format, bg_color=background, asset_loader=asset_loader, bg_image_bytes=bg_img_bytes)

        msg_part = f'<p style="margin:0 0 16px;">{message}</p>' if message else ''
        html_body = (
            '<div style="font-family:Arial,Helvetica,sans-serif;max-width:640px;margin:0 auto;padding:16px;">'
            f'{msg_part}'
            '<p style="color:#666;font-size:13px;margin:12px 0 0;">See the attached flyer.</p>'
            '</div>'
        )

        safe_name = ''.join(c if c.isalnum() or c in ' -_' else '_' for c in t.name)
        png_filename = f'{safe_name}.png'
        recipients = [c.email for c in contacts]
        template_id_val, template_name = t.id, t.name
        db.session.close()

        try:
            sent, failed = send_flyer_bulk_smtp(recipients, subject, html_body, png_bytes, png_filename)
        except RuntimeError as e:
            return jsonify({'error': str(e)}), 500
        except Exception as e:
            return jsonify({'error': f'Could not send: {e}'}), 502

        log_audit('flyer_template_sent', 'flyer_template', template_id_val, template_name, {
            'recipient_count': len(recipients), 'sent': sent, 'failed': failed,
        })
        return jsonify({'sent': sent, 'failed': failed, 'total': len(recipients)})

    @app.route('/api/flyer-templates/<int:template_id>/render', methods=['POST'])
    def render_flyer_template(template_id):
        import base64
        from flyer_render import render_flyer_png, CANVAS_FORMATS

        t = FlyerTemplate.query.get_or_404(template_id)
        data = request.get_json(silent=True) or {}
        elements = data.get('elements') or t.elements or []
        fmt = t.format

        def asset_loader(asset_id):
            try:
                a = FlyerAsset.query.get(int(asset_id))
                return a.data if a else None
            except (TypeError, ValueError):
                return None

        bg = data.get('background', t.background or '#ffffff')
        bg_asset_id = data.get('bg_asset_id', t.bg_asset_id)
        bg_image_bytes = None
        if bg_asset_id:
            a = FlyerAsset.query.get(int(bg_asset_id))
            if a:
                bg_image_bytes = a.data
        png_bytes = render_flyer_png(elements, fmt=fmt, bg_color=bg, asset_loader=asset_loader, bg_image_bytes=bg_image_bytes)
        return jsonify({
            'image': 'data:image/png;base64,' + base64.b64encode(png_bytes).decode(),
            'width':  CANVAS_FORMATS.get(fmt, CANVAS_FORMATS['square'])['rw'],
            'height': CANVAS_FORMATS.get(fmt, CANVAS_FORMATS['square'])['rh'],
        })

    @app.route('/api/flyer-assets', methods=['POST'])
    def upload_flyer_asset():
        f = request.files.get('file')
        if not f or not f.filename:
            return jsonify({'error': 'No file provided.'}), 400
        data = f.read()
        if len(data) > 10 * 1024 * 1024:
            return jsonify({'error': 'File too large (10MB limit).'}), 400
        asset = FlyerAsset(data=data, mimetype=f.mimetype or 'application/octet-stream')
        db.session.add(asset)
        db.session.commit()
        return jsonify({'id': asset.id, 'mimetype': asset.mimetype}), 201

    @app.route('/flyer-builder/assets/<int:asset_id>')
    def serve_flyer_asset(asset_id):
        asset = FlyerAsset.query.get_or_404(asset_id)
        return send_file(
            io.BytesIO(asset.data),
            mimetype=asset.mimetype,
        )

    @app.route('/api/contacts', methods=['GET'])
    def list_contacts():
        q = request.args.get('q', type=str)
        tag = parse_multi_param('tag')
        county = parse_multi_param('county')
        followup = request.args.get('followup', type=str)
        favorites_only = request.args.get('favorites_only', type=str) in ('1', 'true', 'True')
        page = request.args.get('page', default=1, type=int)
        limit = request.args.get('limit', default=25, type=int)

        query = filtered_contacts_query(q=q, tag=tag, county=county, followup=followup, favorites_only=favorites_only)

        total = query.count()
        results = query.order_by(Contact.added.desc()).offset((page - 1) * limit).limit(limit).all()
        contacts = contacts_schema.dump(results)

        # Batched per-page lookup of outreach recency, not per-contact --
        # this page has at most `limit` rows, so one extra query here is
        # cheap regardless of how many contacts exist in total.
        contact_ids = [c.id for c in results]
        if contact_ids:
            rows = (
                db.session.query(Activity.contact_id, Activity.channel, func.max(Activity.contacted_on))
                .filter(Activity.contact_id.in_(contact_ids))
                .group_by(Activity.contact_id, Activity.channel)
                .all()
            )
            last_contacted = {}
            last_emailed = {}
            for cid, channel, latest in rows:
                if latest and (cid not in last_contacted or latest > last_contacted[cid]):
                    last_contacted[cid] = latest
                if channel == 'Email' and latest and (cid not in last_emailed or latest > last_emailed[cid]):
                    last_emailed[cid] = latest
            for c in contacts:
                lc = last_contacted.get(c['id'])
                le = last_emailed.get(c['id'])
                c['last_contacted_on'] = lc.isoformat() if lc else None
                c['last_emailed_on'] = le.isoformat() if le else None

        return jsonify({
            'page': page,
            'limit': limit,
            'total': total,
            'contacts': contacts
        })

    @app.route('/api/contacts/<int:contact_id>', methods=['GET'])
    def get_contact(contact_id):
        c = Contact.query.get_or_404(contact_id)
        return jsonify(contact_schema.dump(c))

    @app.route('/api/contacts/<int:contact_id>', methods=['PUT'])
    def update_contact(contact_id):
        c = Contact.query.get_or_404(contact_id)
        data = request.get_json() or {}
        new_email = (data.get('email') or '').strip() or None
        if 'email' in data and new_email and new_email != c.email:
            conflict = Contact.query.filter(func.lower(Contact.email) == new_email.lower(), Contact.id != contact_id).first()
            if conflict:
                return jsonify({'error': 'email exists', 'id': conflict.id}), 409

        changes = {}
        if 'email' in data and new_email != c.email:
            changes['email'] = {'old': c.email, 'new': new_email}
        for field in ['first_name','last_name','organization','title','phone_office','phone_cell','active','county','notes','tag']:
            if field in data:
                old = getattr(c, field)
                new = data.get(field)
                if old != new:
                    changes[field] = {'old': old, 'new': new}
                setattr(c, field, new)
        if 'email' in data:
            c.email = new_email
        if 'lists' in data:
            new_lists = data.get('lists') or []
            if (c.lists or []) != new_lists:
                changes['lists'] = {'old': c.lists or [], 'new': new_lists}
            c.lists = new_lists
        if 'data_complete' in data:
            new_dc = bool(data.get('data_complete'))
            if bool(c.data_complete) != new_dc:
                changes['data_complete'] = {'old': bool(c.data_complete), 'new': new_dc}
            c.data_complete = new_dc
        db.session.add(c)
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': 'database error', 'details': str(e)}), 500
        if changes:
            label = f"{c.first_name or ''} {c.last_name or ''}".strip() or c.email or f'#{c.id}'
            log_audit('contact_updated', 'contact', c.id, label, changes)
        return jsonify(contact_schema.dump(c))

    @app.route('/api/contacts/<int:contact_id>', methods=['DELETE'])
    def delete_contact(contact_id):
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403
        c = Contact.query.get_or_404(contact_id)
        label = f"{c.first_name or ''} {c.last_name or ''}".strip() or c.email or f'#{c.id}'
        # No ON DELETE CASCADE on activities.contact_id -- deleting the
        # contact's own outreach log entries first avoids a FK violation
        # on Postgres (SQLite doesn't enforce this by default, but
        # Postgres does).
        Activity.query.filter_by(contact_id=c.id).delete()
        db.session.delete(c)
        db.session.commit()
        log_audit('contact_deleted', 'contact', contact_id, label)
        return jsonify({'deleted': True})

    @app.route('/api/contacts/<int:contact_id>/favorite', methods=['PUT'])
    def toggle_contact_favorite(contact_id):
        # Deliberately not run through update_contact()'s change-diffing --
        # starring something is a personal/team organizing action, not a
        # data edit worth cluttering the Audit Log over.
        c = Contact.query.get_or_404(contact_id)
        data = request.get_json(silent=True) or {}
        c.is_favorite = bool(data.get('is_favorite'))
        db.session.commit()
        return jsonify({'id': c.id, 'is_favorite': c.is_favorite})

    @app.route('/api/contacts', methods=['POST'])
    def create_contact():
        data = request.get_json() or {}
        email = (data.get('email') or '').strip() or None
        force = bool(data.get('force_create'))
        if email:
            existing = Contact.query.filter(func.lower(Contact.email) == email.lower()).first()
            if existing:
                return jsonify({'error': 'email exists', 'id': existing.id}), 409
        if not force:
            first_name = (data.get('first_name') or '').strip()
            last_name = (data.get('last_name') or '').strip()
            organization = (data.get('organization') or '').strip()
            if first_name and last_name and organization:
                possible_dup = Contact.query.filter(
                    func.lower(Contact.first_name) == first_name.lower(),
                    func.lower(Contact.last_name) == last_name.lower(),
                    func.lower(Contact.organization) == organization.lower(),
                ).first()
                if possible_dup:
                    return jsonify({
                        'warning': 'possible_duplicate',
                        'id': possible_dup.id,
                        'message': f'A contact named {first_name} {last_name} at {organization} already exists. Add anyway?',
                    }), 409
        c = Contact(
            email=email,
            first_name=data.get('first_name') or None,
            last_name=data.get('last_name') or None,
            organization=data.get('organization') or None,
            title=data.get('title') or None,
            phone_office=data.get('phone_office') or None,
            phone_cell=data.get('phone_cell') or None,
            active=data.get('active') or None,
            lists=data.get('lists') or [],
            county=data.get('county') or None,
            notes=data.get('notes') or None,
            tag=data.get('tag') or None,
            data_complete=bool(data.get('data_complete')),
        )
        db.session.add(c)
        db.session.commit()
        label = f"{c.first_name or ''} {c.last_name or ''}".strip() or c.email or f'#{c.id}'
        log_audit('contact_created', 'contact', c.id, label)
        return jsonify(contact_schema.dump(c)), 201

    @app.route('/api/contacts/<int:contact_id>/activity', methods=['GET'])
    def list_contact_activity(contact_id):
        Contact.query.get_or_404(contact_id)
        rows = Activity.query.filter_by(contact_id=contact_id) \
            .order_by(Activity.contacted_on.desc(), Activity.created_at.desc()).all()
        return jsonify({'activity': [a.to_dict() for a in rows]})

    @app.route('/api/contacts/<int:contact_id>/activity', methods=['POST'])
    def create_contact_activity(contact_id):
        c = Contact.query.get_or_404(contact_id)
        data = request.get_json(silent=True) or {}
        summary = (data.get('summary') or '').strip()
        if not summary:
            return jsonify({'error': 'summary is required'}), 400
        a = Activity(
            contact_id=c.id,
            organization=c.organization,
            employee_name=current_user.display_name,
            channel=data.get('channel') or None,
            summary=summary,
            contacted_on=date.fromisoformat(data['contacted_on']) if data.get('contacted_on') else date.today(),
        )
        db.session.add(a)
        db.session.commit()
        return jsonify(a.to_dict()), 201

    @app.route('/api/organizations/<organization>/activity', methods=['GET'])
    def list_org_activity(organization):
        rows = Activity.query.filter(func.lower(Activity.organization) == organization.lower()) \
            .order_by(Activity.contacted_on.desc(), Activity.created_at.desc()).all()
        return jsonify({'activity': [a.to_dict() for a in rows]})

    @app.route('/api/organizations/<organization>/activity', methods=['POST'])
    def create_org_activity(organization):
        data = request.get_json(silent=True) or {}
        summary = (data.get('summary') or '').strip()
        if not summary:
            return jsonify({'error': 'summary is required'}), 400
        a = Activity(
            contact_id=None,
            organization=organization,
            employee_name=current_user.display_name,
            channel=data.get('channel') or None,
            summary=summary,
            contacted_on=date.fromisoformat(data['contacted_on']) if data.get('contacted_on') else date.today(),
        )
        db.session.add(a)
        db.session.commit()
        return jsonify(a.to_dict()), 201

    @app.route('/api/activity/<int:activity_id>', methods=['DELETE'])
    def delete_activity(activity_id):
        a = Activity.query.get_or_404(activity_id)
        db.session.delete(a)
        db.session.commit()
        return jsonify({'ok': True})

    @app.route('/api/stats', methods=['GET'])
    def stats():
        total = Contact.query.count()
        incomplete = Contact.query.filter(contact_incomplete_clause()).count()
        organizations = OutreachOrg.query.count()
        complete_pct = round(100 * (total - incomplete) / total) if total else 0
        per_tag = db.session.query(Contact.tag, func.count(Contact.id)).group_by(Contact.tag).all()
        per_county = db.session.query(Contact.county, func.count(Contact.id)).group_by(Contact.county).all()
        return jsonify({
            'total': total,
            'incomplete': incomplete,
            'organizations': organizations,
            'complete_pct': complete_pct,
            'by_tag': {k if k else '': v for k, v in per_tag},
            'by_county': {k if k else '': v for k, v in per_county}
        })

    @app.route('/api/tags', methods=['GET'])
    def tags():
        tags = [t[0] for t in db.session.query(Contact.tag).distinct().all()]
        return jsonify(sorted([t for t in tags if t]))

    @app.route('/api/categories', methods=['GET'])
    def categories():
        # Categories map to tags with counts
        rows = db.session.query(Contact.tag, func.count(Contact.id)).group_by(Contact.tag).all()
        return jsonify([{'tag': r[0], 'count': r[1]} for r in rows])

    @app.route('/api/section-categories', methods=['GET'])
    def section_categories():
        # Categories for the Sections/outreach-checklist page (its own tag set,
        # separate from Contact.tag since the two sheets label categories differently).
        tags = [t[0] for t in db.session.query(OutreachOrg.tag).distinct().all()]
        return jsonify(sorted([t for t in tags if t]))

    @app.route('/api/section-stats', methods=['GET'])
    def section_stats():
        total = OutreachOrg.query.count()
        contact_orgs = set(
            (o[0] or '').lower() for o in db.session.query(Contact.organization).filter(Contact.organization != None).all()
        )
        org_names = db.session.query(OutreachOrg.organization).all()
        no_contact = sum(1 for o in org_names if o[0].lower() not in contact_orgs)
        return jsonify({'total': total, 'no_contact': no_contact})

    @app.route('/api/counties', methods=['GET'])
    def counties():
        # Contact.county can hold several comma-separated counties in one
        # string (e.g. "Dallas, Tarrant, Collin") -- split those apart so
        # each county name is offered exactly once in the filter, instead
        # of every distinct combination showing up as its own option.
        rows = db.session.query(Contact.county).filter(Contact.county.isnot(None), Contact.county != '').distinct().all()
        names = set()
        for (val,) in rows:
            for part in (val or '').split(','):
                part = part.strip()
                if part:
                    names.add(part)
        return jsonify(sorted(names))

    @app.route('/api/sections', methods=['GET'])
    def sections():
        """Organizations view: the outreach checklist (OutreachOrg: category,
        org, last-touched date, notes) cross-referenced against the live
        Contact table by organization name, returned as a flat paginated list
        (same shape as /api/contacts) with every matching contact at that
        organization -- not just one "primary" contact -- so coworkers who
        share an organization show up together.

        Supports optional filters: q (text), tag (category), county, page, limit.
        """
        q = request.args.get('q', type=str)
        tag_filter = parse_multi_param('tag')
        county = parse_multi_param('county')
        page = request.args.get('page', default=1, type=int)
        limit = request.args.get('limit', default=25, type=int)

        org_query = OutreachOrg.query
        if tag_filter:
            org_query = org_query.filter(OutreachOrg.tag.in_(tag_filter))
        if q:
            like = f"%{q}%"
            org_query = org_query.filter(or_(OutreachOrg.organization.ilike(like), OutreachOrg.tag.ilike(like), OutreachOrg.notes.ilike(like)))

        rows = org_query.order_by(OutreachOrg.tag, OutreachOrg.organization).all()

        # Fetch every relevant contact in one query and group by lowercased
        # organization name in Python, instead of issuing a query per
        # organization -- that N+1 pattern was the main reason this endpoint
        # felt slow/unresponsive with ~280 organizations.
        contacts_query = Contact.query
        county_clause = county_filter_clause(county)
        if county_clause is not None:
            contacts_query = contacts_query.filter(county_clause)
        contacts_by_org = {}
        for c in contacts_query.order_by(Contact.first_name, Contact.last_name).all():
            if not c.organization:
                continue
            contacts_by_org.setdefault(c.organization.lower(), []).append(c)

        items = []
        for org_row in rows:
            org = org_row.organization
            org_contacts = contacts_by_org.get(org.lower(), [])
            if county and not org_contacts:
                # no contacts in the requested county for this org; skip it
                continue
            items.append({
                'tag': org_row.tag or 'Other',
                'organization': org,
                'contact_count': len(org_contacts),
                'latest_updated': org_row.updated,
                'contacts': [{
                    'id': c.id,
                    'name': f"{c.first_name or ''} {c.last_name or ''}".strip(),
                    'title': c.title or '',
                    'email': c.email or '',
                    'phone_cell': c.phone_cell or '',
                    'phone_office': c.phone_office or '',
                } for c in org_contacts],
                'notes': org_row.notes or ''
            })

        # sort by last-touched date desc (untouched orgs sort last)
        items.sort(key=lambda it: it['latest_updated'] or date(1970, 1, 1), reverse=True)

        total = len(items)
        start = (page - 1) * limit
        page_items = items[start:start + limit]
        for it in page_items:
            it['latest_updated'] = it['latest_updated'].isoformat() if it['latest_updated'] else None

        return jsonify({'page': page, 'limit': limit, 'total': total, 'organizations': page_items})

    @app.route('/api/upload', methods=['POST'])
    def upload():
        # Accepts a CSV (People only) or an Excel workbook -- a workbook
        # can have a People tab and a separate Organizations tab (matching
        # the original spreadsheet's shape); each sheet is classified by
        # its columns and upserted with the matching importer.
        if not current_user.is_admin:
            return jsonify({'error': 'admin only'}), 403
        if 'file' not in request.files:
            return jsonify({'error': 'file is required (form field `file`)'}), 400
        f = request.files['file']
        try:
            sheets = read_uploaded_file(f)
        except Exception as e:
            return jsonify({'error': str(e)}), 400

        contacts_result = {'inserted': 0, 'updated': 0, 'skipped': 0}
        orgs_result = {'inserted': 0, 'updated': 0, 'skipped': 0}

        for df in sheets.values():
            if df is None or df.empty:
                continue
            if looks_like_contacts_sheet(df):
                _import_contacts(df, contacts_result)
            elif looks_like_orgs_sheet(df):
                _import_orgs(df, orgs_result)
            # Sheets that match neither shape (e.g. an instructions tab)
            # are silently ignored rather than guessed at.

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': 'database error', 'details': str(e)}), 500

        log_audit('spreadsheet_sync', 'sync', None, f.filename, {
            'contacts': contacts_result,
            'organizations': orgs_result,
        })
        return jsonify({'contacts': contacts_result, 'organizations': orgs_result})

    @app.route('/api/export', methods=['GET'])
    def export():
        q = request.args.get('q', type=str)
        tag = parse_multi_param('tag')
        org_tag = parse_multi_param('org_tag')
        county = parse_multi_param('county')
        followup = request.args.get('followup', type=str)
        favorites_only = request.args.get('favorites_only', type=str) in ('1', 'true', 'True')
        contact_id = request.args.get('id', type=int)
        rows = filtered_contacts_query(q=q, tag=tag, org_tag=org_tag, county=county, contact_id=contact_id, followup=followup, favorites_only=favorites_only).order_by(Contact.added.desc()).all()

        # stream CSV
        si = io.StringIO()
        writer = csv.writer(si)
        header = ['id','tag','organization','first_name','last_name','title','phone_office','phone_cell','email','added','active','lists','county','notes','data_complete']
        writer.writerow(header)
        for r in rows:
            writer.writerow([
                r.id, r.tag, r.organization, r.first_name, r.last_name, r.title,
                r.phone_office, r.phone_cell, r.email, r.added.isoformat() if r.added else '',
                r.active, (','.join(r.lists) if r.lists else ''), r.county, r.notes, int(bool(r.data_complete))
            ])
        si.seek(0)
        return send_file(io.BytesIO(si.getvalue().encode('utf-8')), mimetype='text/csv', as_attachment=True, download_name='contacts_export.csv')

    @app.route('/api/export/emails', methods=['GET'])
    def export_emails():
        """Flat, de-duplicated list of email addresses for the current filter,
        meant for pasting into the BCC field of a mass email."""
        q = request.args.get('q', type=str)
        tag = parse_multi_param('tag')
        org_tag = parse_multi_param('org_tag')
        county = parse_multi_param('county')
        followup = request.args.get('followup', type=str)
        favorites_only = request.args.get('favorites_only', type=str) in ('1', 'true', 'True')
        rows = filtered_contacts_query(q=q, tag=tag, org_tag=org_tag, county=county, followup=followup, favorites_only=favorites_only).all()

        emails = []
        seen = set()
        for r in rows:
            if not r.email:
                continue
            # some legacy records store multiple addresses separated by '|'
            for part in r.email.split('|'):
                addr = part.strip()
                if addr and addr.lower() not in seen:
                    seen.add(addr.lower())
                    emails.append(addr)

        return jsonify({'count': len(emails), 'emails': emails, 'joined': ', '.join(emails)})

    @app.route('/api/export/docx', methods=['GET'])
    def export_docx():
        from docx import Document

        q = request.args.get('q', type=str)
        tag = parse_multi_param('tag')
        org_tag = parse_multi_param('org_tag')
        county = parse_multi_param('county')
        followup = request.args.get('followup', type=str)
        favorites_only = request.args.get('favorites_only', type=str) in ('1', 'true', 'True')
        rows = filtered_contacts_query(q=q, tag=tag, org_tag=org_tag, county=county, followup=followup, favorites_only=favorites_only).order_by(Contact.organization, Contact.last_name).all()

        doc = Document()
        title = ', '.join(tag) if tag else (', '.join(org_tag) if org_tag else 'Contacts')
        doc.add_heading(title, level=1)
        doc.add_paragraph(f'{len(rows)} contact(s)')

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = 'Name', 'Title', 'Organization', 'Email', 'Phone'
        for r in rows:
            cells = table.add_row().cells
            cells[0].text = f"{r.first_name or ''} {r.last_name or ''}".strip()
            cells[1].text = r.title or ''
            cells[2].text = r.organization or ''
            cells[3].text = r.email or ''
            cells[4].text = r.phone_office or r.phone_cell or ''

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_name = (tag or org_tag or 'contacts').replace('/', '-').replace(' ', '_')
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          as_attachment=True, download_name=f'{safe_name}_export.docx')

    @app.route('/api/draft-email', methods=['POST'])
    def draft_email():
        import anthropic

        data = request.get_json(silent=True) or {}
        prompt = (data.get('prompt') or '').strip()
        if not prompt:
            return jsonify({'error': 'Describe the email you want to draft.'}), 400

        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            return jsonify({'error': 'ANTHROPIC_API_KEY is not configured on the server.'}), 500

        q = data.get('q')
        tag = data.get('tag')
        org_tag = data.get('org_tag')
        county = data.get('county')
        followup = data.get('followup')
        favorites_only = bool(data.get('favorites_only'))
        rows = filtered_contacts_query(q=q, tag=tag, org_tag=org_tag, county=county, followup=followup, favorites_only=favorites_only).all()
        orgs = sorted({r.organization for r in rows if r.organization})

        context_lines = [f"{len(rows)} recipient(s) in this group."]
        if tag:
            context_lines.append(f'Category/tag filter: {tag}.')
        if org_tag:
            context_lines.append(f'Organization category filter: {org_tag}.')
        if county:
            context_lines.append(f'County filter: {county}.')
        if q:
            context_lines.append(f'Search filter: "{q}".')
        if followup == 'never':
            context_lines.append("This group has never been contacted before -- this would be a first outreach, not a check-in.")
        elif followup:
            context_lines.append(f"This group hasn't been contacted in {followup}+ days -- consider a warmer, re-engaging tone rather than a routine update.")
        if orgs:
            sample = ', '.join(orgs[:8])
            context_lines.append(f'Sample organizations: {sample}{"..." if len(orgs) > 8 else "."}')

        case_study_block = ''
        case_study_id = data.get('case_study_id')
        if case_study_id:
            case_study = CaseStudy.query.get(case_study_id)
            if case_study:
                cs_lines = [f"Title: {case_study.title}"]
                if case_study.client:
                    cs_lines.append(f"Client: {case_study.client}")
                if case_study.sector:
                    cs_lines.append(f"Sector: {case_study.sector}")
                cs_text = '\n'.join(filter(None, [case_study.challenges, case_study.solution, case_study.results])) or case_study.extracted_text or ''
                if cs_text:
                    cs_lines.append(cs_text[:3000])
                case_study_block = (
                    "\n\nRelevant past case study you may draw on as a proof point if it "
                    "fits naturally -- use only what's given below, don't invent extra "
                    "detail beyond it:\n" + "\n".join(cs_lines)
                )

        system = (
            "You draft outreach emails for JBJ Management, sent to community contacts "
            "(elected officials, organizations, clergy, chambers of commerce, etc). Write "
            "a complete, professional but warm email with a subject line and body, tailored "
            "to the recipient group described. Use \"[Name]\" as a placeholder for the "
            "individual recipient's name. Do not invent specific facts (dates, addresses, "
            "times) the user didn't provide -- use a placeholder like [DATE] or [LOCATION] "
            "instead. Output only the subject line and email body, no commentary."
        )
        user_message = "Recipient group:\n" + "\n".join(context_lines) + case_study_block + f"\n\nEmail request: {prompt}"

        try:
            client = anthropic.Anthropic(api_key=api_key)
            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1024,
                system=system,
                messages=[{"role": "user", "content": user_message}],
            )
        except anthropic.APIStatusError as e:
            return jsonify({'error': f'Claude API error: {e.message}'}), 502
        except Exception as e:
            return jsonify({'error': str(e)}), 502

        draft = next((b.text for b in response.content if b.type == 'text'), '')
        return jsonify({'draft': draft, 'recipient_count': len(rows)})

    @app.route('/api/generate-flyer', methods=['POST'])
    def generate_flyer():
        import json
        import base64
        import anthropic
        import openai
        from PIL import Image, ImageDraw, ImageFont

        data = request.get_json(silent=True) or {}
        prompt = (data.get('prompt') or '').strip()
        fmt = data.get('format') if data.get('format') in _valid_flyer_formats() else 'square'
        if not prompt:
            return jsonify({'error': 'Describe what the post or flyer is about.'}), 400

        anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
        openai_key = os.environ.get('OPENAI_API_KEY')
        if not anthropic_key:
            return jsonify({'error': 'ANTHROPIC_API_KEY is not configured on the server.'}), 500
        if not openai_key:
            return jsonify({'error': 'OPENAI_API_KEY is not configured on the server.'}), 500

        width, height = (1024, 1536) if fmt == 'portrait' else (1024, 1024)

        copy_system = (
            "You write short marketing copy for a single social media post or printed "
            "flyer image for JBJ Management, a community/government-relations firm. "
            'Respond with ONLY valid JSON, no commentary or markdown fences, in this '
            'exact shape: {"headline": "...", "body": "..."}. The headline must be 3-7 '
            "words, punchy, no ending period. The body must be 1-2 short sentences, 25 "
            "words or fewer. Do not invent specific dates, addresses, or prices the user "
            "didn't provide -- use a placeholder like [DATE] if one seems needed."
        )
        raw_text = ''
        try:
            claude = anthropic.Anthropic(api_key=anthropic_key)
            copy_response = claude.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=200,
                system=copy_system,
                messages=[{"role": "user", "content": prompt}],
            )
            raw_text = next((b.text for b in copy_response.content if b.type == 'text'), '{}').strip()
            if raw_text.startswith('```'):
                raw_text = raw_text.strip('`')
                if raw_text.startswith('json'):
                    raw_text = raw_text[4:]
            parsed = json.loads(raw_text)
            headline = (parsed.get('headline') or '').strip() or 'JBJ Management'
            body = (parsed.get('body') or '').strip()
        except anthropic.APIStatusError as e:
            return jsonify({'error': f'Claude API error: {e.message}'}), 502
        except Exception:
            headline = raw_text[:60] or 'JBJ Management'
            body = ''

        visual_prompt = (
            f"A clean, professional background graphic for a "
            f"{'printed flyer' if fmt == 'portrait' else 'social media post'} about: "
            f"{prompt}. Style: modern, minimal, photo-realistic or tasteful abstract "
            "design. Color palette: deep maroon red, black, and soft dusty rose accents "
            "on white or light background. Leave clear open, uncluttered space in the "
            "lower third of the image for text to be added afterward. Absolutely no "
            "text, words, letters, or numbers anywhere in the image."
        )
        try:
            oai_client = openai.OpenAI(api_key=openai_key)
            img_response = oai_client.images.generate(
                model="gpt-image-1-mini",
                prompt=visual_prompt,
                size=f"{width}x{height}",
                quality="low",
            )
            img_bytes = base64.b64decode(img_response.data[0].b64_json)
        except openai.APIStatusError as e:
            return jsonify({'error': f'OpenAI image API error: {e.message}'}), 502
        except Exception as e:
            return jsonify({'error': str(e)}), 502

        base_img = Image.open(io.BytesIO(img_bytes)).convert('RGBA')
        if base_img.size != (width, height):
            base_img = base_img.resize((width, height))
        overlay = Image.new('RGBA', base_img.size, (0, 0, 0, 0))
        draw = ImageDraw.Draw(overlay)

        fonts_dir = os.path.join(os.path.dirname(__file__), 'static', 'fonts')
        headline_font = ImageFont.truetype(
            os.path.join(fonts_dir, 'ArchivoBlack-Regular.ttf'),
            56 if fmt == 'portrait' else 48,
        )
        body_font = ImageFont.truetype(os.path.join(fonts_dir, 'Inter-Variable.ttf'), 28)
        try:
            body_font.set_variation_by_name('Regular')
        except Exception:
            pass

        margin = 60
        max_width = width - margin * 2

        def wrap_text(text, font, max_w):
            lines, cur = [], ''
            for word in text.split():
                trial = (cur + ' ' + word).strip()
                if draw.textlength(trial, font=font) <= max_w:
                    cur = trial
                else:
                    if cur:
                        lines.append(cur)
                    cur = word
            if cur:
                lines.append(cur)
            return lines

        headline_lines = wrap_text(headline, headline_font, max_width)
        body_lines = wrap_text(body, body_font, max_width) if body else []
        line_gap = 10
        headline_h = len(headline_lines) * (headline_font.size + line_gap)
        body_h = len(body_lines) * (body_font.size + 8) if body_lines else 0
        band_height = min(height, headline_h + body_h + 80)
        band_top = height - band_height

        draw.rectangle([0, band_top, width, height], fill=(20, 0, 2, 175))

        y = band_top + 36
        for line in headline_lines:
            draw.text((margin, y), line, font=headline_font, fill=(255, 255, 255, 255))
            y += headline_font.size + line_gap
        y += 8
        for line in body_lines:
            draw.text((margin, y), line, font=body_font, fill=(240, 240, 240, 255))
            y += body_font.size + 8

        final = Image.alpha_composite(base_img, overlay).convert('RGB')
        buf = io.BytesIO()
        final.save(buf, format='PNG')
        final_b64 = base64.b64encode(buf.getvalue()).decode('ascii')

        return jsonify({
            'image': f'data:image/png;base64,{final_b64}',
            'headline': headline,
            'body': body,
        })

    # ------------------------------------------------------------------ #
    # Send Campaign                                                        #
    # ------------------------------------------------------------------ #

    def _text_to_html(text):
        """Plain-text body → minimal HTML paragraphs suitable for email."""
        paragraphs = text.split('\n\n')
        parts = []
        for p in paragraphs:
            p = p.strip()
            if p:
                parts.append(
                    f'<p style="margin:0 0 12px;line-height:1.6;">{p.replace(chr(10),"<br>")}</p>'
                )
        return (
            '<div style="font-family:Arial,Helvetica,sans-serif;max-width:600px;'
            'margin:0 auto;padding:16px;">'
            + ''.join(parts)
            + '</div>'
        )

    @app.route('/api/campaign/preview', methods=['GET'])
    def campaign_preview():
        q        = request.args.get('q')
        tag      = request.args.get('tag')
        county   = request.args.get('county')
        followup = request.args.get('followup')
        favorites_only = request.args.get('favorites_only') == 'true'
        contacts = filtered_contacts_query(
            q=q, tag=tag, county=county, followup=followup, favorites_only=favorites_only
        ).filter(
            Contact.email.isnot(None),
            Contact.email != '',
            Contact.unsubscribed == False,
        ).all()
        sample = [
            f"{(c.first_name or '')} {(c.last_name or '')}".strip() or c.email
            for c in contacts[:5]
        ]
        return jsonify({'recipient_count': len(contacts), 'sample': sample})

    @app.route('/api/campaign/send', methods=['POST'])
    def campaign_send():
        data    = request.get_json() or {}
        subject = (data.get('subject') or '').strip()
        body    = (data.get('body') or '').strip()
        if not subject:
            return jsonify({'error': 'Subject is required.'}), 400
        if not body:
            return jsonify({'error': 'Email body is required.'}), 400

        q        = data.get('q')
        tag      = data.get('tag')
        county   = data.get('county')
        followup = data.get('followup')
        favorites_only = bool(data.get('favorites_only'))

        contacts = filtered_contacts_query(
            q=q, tag=tag, county=county, followup=followup, favorites_only=favorites_only
        ).filter(
            Contact.email.isnot(None),
            Contact.email != '',
            Contact.unsubscribed == False,
        ).all()

        if not contacts:
            return jsonify({'error': 'No contacts with email addresses match this filter.'}), 400

        html_body  = _text_to_html(body)
        recipients = [c.email for c in contacts]
        contact_ids = [c.id for c in contacts]
        sent_by    = current_user.display_name

        # Snapshot the send record before releasing the DB connection
        send_rec = EmailSend(
            sent_by_name=sent_by,
            subject=subject,
            filter_snapshot={'q': q, 'tag': tag, 'county': county, 'followup': followup},
            recipient_count=len(contacts),
            status='sending',
            started_at=datetime.utcnow(),
        )
        db.session.add(send_rec)
        db.session.commit()
        send_id = send_rec.id

        # Release connection before the slow network call (same pattern as
        # the email-template send route -- avoids idle-connection SSL errors).
        db.session.close()

        # Use tracked per-recipient sends when SendGrid API key is available
        sg_key = os.environ.get('SENDGRID_API_KEY') or (
            os.environ.get('SMTP_PASSWORD', '').startswith('SG.') and os.environ.get('SMTP_PASSWORD')
        ) or None
        recipients_with_ids = [{'email': c.email, 'contact_id': c.id} for c in contacts]
        try:
            if sg_key:
                sent, failed = _sendgrid_campaign_tracked(sg_key, recipients_with_ids, subject, html_body, send_id)
            else:
                sent, failed = send_flyer_bulk_smtp(recipients, subject, html_body)
        except RuntimeError as e:
            return jsonify({'error': str(e)}), 500
        except Exception as e:
            return jsonify({'error': f'Could not send: {e}'}), 502

        # Update the send record status and log one outreach activity per contact
        send_rec2 = EmailSend.query.get(send_id)
        if send_rec2:
            send_rec2.sent_count = sent
            send_rec2.failed_count = failed
            send_rec2.status = 'completed'
            send_rec2.completed_at = datetime.utcnow()
            db.session.commit()

        for cid in contact_ids:
            act = Activity(
                contact_id=cid,
                employee_name=sent_by,
                channel='Email',
                summary=f'Campaign email sent: {subject}',
                contacted_on=date.today(),
            )
            db.session.add(act)
        db.session.commit()

        log_audit('campaign_sent', 'campaign', send_id, subject, {
            'recipient_count': len(recipients), 'sent': sent, 'failed': failed,
        })
        return jsonify({'sent': sent, 'failed': failed, 'total': len(recipients)})

    # ------------------------------------------------------------------ #
    # Tasks                                                                #
    # ------------------------------------------------------------------ #

    def _task_urgency(task, today):
        if task.due_date is None:
            return 'no_date'
        if task.due_date < today:
            return 'overdue'
        if task.due_date == today:
            return 'today'
        return 'upcoming'

    @app.route('/api/tasks/count', methods=['GET'])
    def task_count():
        today = date.today()
        count = Task.query.filter(
            Task.created_by_id == current_user.id,
            Task.completed == False,
            Task.due_date != None,
            Task.due_date <= today,
        ).count()
        return jsonify({'count': count})

    @app.route('/api/tasks', methods=['GET'])
    def list_tasks():
        today = date.today()
        show_completed = request.args.get('completed', 'false').lower() == 'true'
        query = Task.query.filter_by(completed=show_completed, created_by_id=current_user.id)

        if show_completed:
            tasks = query.order_by(Task.completed_at.desc()).all()
        else:
            tasks = query.all()

            def sort_key(t):
                urgency = _task_urgency(t, today)
                order = {'overdue': 0, 'today': 1, 'upcoming': 2, 'no_date': 3}
                d = t.due_date or date.max
                return (order[urgency], d)

            tasks.sort(key=sort_key)

        result = []
        for t in tasks:
            d = t.to_dict()
            d['urgency'] = _task_urgency(t, today) if not t.completed else 'done'
            result.append(d)
        return jsonify({'tasks': result})

    @app.route('/api/contacts/<int:contact_id>/tasks', methods=['GET'])
    def contact_tasks(contact_id):
        Contact.query.get_or_404(contact_id)
        today = date.today()
        tasks = Task.query.filter_by(contact_id=contact_id).order_by(
            Task.completed.asc(), Task.due_date.asc()
        ).all()
        result = []
        for t in tasks:
            d = t.to_dict()
            d['urgency'] = _task_urgency(t, today) if not t.completed else 'done'
            result.append(d)
        return jsonify({'tasks': result})

    @app.route('/api/tasks', methods=['POST'])
    def create_task():
        data = request.get_json() or {}
        title = (data.get('title') or '').strip()
        if not title:
            return jsonify({'error': 'title required'}), 400
        due_date = None
        if data.get('due_date'):
            try:
                due_date = date.fromisoformat(data['due_date'])
            except ValueError:
                return jsonify({'error': 'invalid due_date'}), 400
        task = Task(
            contact_id=data.get('contact_id') or None,
            title=title,
            due_date=due_date,
            notes=(data.get('notes') or '').strip() or None,
            created_by_id=current_user.id,
        )
        db.session.add(task)
        db.session.commit()
        d = task.to_dict()
        today = date.today()
        d['urgency'] = _task_urgency(task, today)
        return jsonify(d), 201

    @app.route('/api/tasks/<int:task_id>', methods=['PATCH'])
    def update_task(task_id):
        task = Task.query.get_or_404(task_id)
        data = request.get_json() or {}
        if 'completed' in data:
            task.completed = bool(data['completed'])
            task.completed_at = datetime.utcnow() if task.completed else None
        if 'title' in data:
            title = (data['title'] or '').strip()
            if title:
                task.title = title
        if 'due_date' in data:
            if data['due_date']:
                try:
                    task.due_date = date.fromisoformat(data['due_date'])
                except ValueError:
                    return jsonify({'error': 'invalid due_date'}), 400
            else:
                task.due_date = None
        if 'notes' in data:
            task.notes = (data['notes'] or '').strip() or None
        db.session.commit()
        d = task.to_dict()
        today = date.today()
        d['urgency'] = _task_urgency(task, today) if not task.completed else 'done'
        return jsonify(d)

    @app.route('/api/tasks/<int:task_id>', methods=['DELETE'])
    def delete_task(task_id):
        task = Task.query.get_or_404(task_id)
        db.session.delete(task)
        db.session.commit()
        return '', 204

    # Pipeline                                                               #

    PIPELINE_STAGES = ['Lead', 'Engaged', 'Proposal', 'Client', 'Inactive']

    @app.route('/api/pipeline', methods=['GET'])
    @login_required
    def get_pipeline():
        contacts = Contact.query.filter(
            Contact.pipeline_stage.isnot(None)
        ).order_by(Contact.first_name, Contact.last_name).all()
        grouped = {s: [] for s in PIPELINE_STAGES}
        for c in contacts:
            stage = c.pipeline_stage
            if stage in grouped:
                grouped[stage].append({
                    'id': c.id,
                    'name': ' '.join(p for p in [c.first_name or '', c.last_name or ''] if p).strip() or c.organization or '(no name)',
                    'organization': c.organization,
                    'tag': c.tag,
                    'email': c.email,
                    'pipeline_stage': stage,
                })
        return jsonify({'stages': PIPELINE_STAGES, 'contacts': grouped})

    @app.route('/api/pipeline/<int:contact_id>', methods=['PATCH'])
    @login_required
    def update_pipeline_stage(contact_id):
        c = Contact.query.get_or_404(contact_id)
        data = request.get_json(force=True)
        stage = data.get('pipeline_stage')
        if stage is not None and stage not in PIPELINE_STAGES and stage != '':
            return jsonify({'error': 'invalid stage'}), 400
        c.pipeline_stage = stage if stage else None
        db.session.commit()
        log_audit('pipeline_stage_changed', 'contact', c.id,
                  f"{c.first_name or ''} {c.last_name or ''}".strip() or c.organization,
                  {'stage': stage})
        # Auto-enroll in active sequences triggered by this stage
        if stage:
            _enroll_contact_in_sequences(c.id, stage)
        return jsonify({'ok': True, 'pipeline_stage': c.pipeline_stage})

    # Email Tracking                                                         #

    with app.app_context():
        db.create_all()  # creates email_events table if missing

    @app.route('/webhooks/sendgrid', methods=['POST'])
    def sendgrid_webhook():
        """Receives open/click/bounce events from SendGrid Event Webhook.
        Configure in SendGrid: Settings → Mail Settings → Event Webhook
        URL: https://your-domain.com/webhooks/sendgrid
        Events to enable: Open, Click, Delivered, Bounce"""
        events = request.get_json(force=True, silent=True) or []
        if not isinstance(events, list):
            events = [events]
        for ev in events:
            event_type = ev.get('event', '')
            if event_type not in ('open', 'click', 'delivered', 'bounce', 'unsubscribe'):
                continue
            email = ev.get('email', '')
            custom = ev.get('custom_args') or {}
            send_id    = int(custom['send_id'])    if custom.get('send_id')    else None
            contact_id = int(custom['contact_id']) if custom.get('contact_id') else None
            ts = ev.get('timestamp')
            occurred_at = datetime.utcfromtimestamp(int(ts)) if ts else datetime.utcnow()
            record = EmailEvent(
                send_id=send_id,
                contact_id=contact_id,
                email=email,
                event_type=event_type,
                url=ev.get('url'),
                sg_message_id=ev.get('sg_message_id'),
                occurred_at=occurred_at,
            )
            db.session.add(record)
            # Auto-unsubscribe on unsubscribe event
            if event_type == 'unsubscribe' and email:
                c = Contact.query.filter_by(email=email).first()
                if c and not c.unsubscribed:
                    c.unsubscribed = True
        db.session.commit()
        return '', 204

    @app.route('/api/contacts/<int:contact_id>/email-events', methods=['GET'])
    @login_required
    def contact_email_events(contact_id):
        events = (EmailEvent.query
                  .filter_by(contact_id=contact_id)
                  .order_by(EmailEvent.occurred_at.desc())
                  .limit(50).all())
        summary = {}
        for e in events:
            summary[e.event_type] = summary.get(e.event_type, 0) + 1
        return jsonify({'events': [e.to_dict() for e in events], 'summary': summary})

    @app.route('/api/email-sends', methods=['GET'])
    @login_required
    def list_email_sends():
        page     = request.args.get('page', 1, type=int)
        per_page = 20
        query    = EmailSend.query.order_by(EmailSend.started_at.desc())
        total    = query.count()
        sends    = query.offset((page - 1) * per_page).limit(per_page).all()
        # Attach open/click counts from email_events
        send_ids = [s.id for s in sends]
        from sqlalchemy import func as sqlfunc
        counts = {}
        if send_ids:
            rows = (db.session.query(
                        EmailEvent.send_id,
                        EmailEvent.event_type,
                        sqlfunc.count(EmailEvent.id).label('n'),
                        sqlfunc.count(sqlfunc.distinct(EmailEvent.contact_id)).label('unique_n'),
                    )
                    .filter(EmailEvent.send_id.in_(send_ids))
                    .group_by(EmailEvent.send_id, EmailEvent.event_type)
                    .all())
            for row in rows:
                counts.setdefault(row.send_id, {})[row.event_type] = {
                    'total': row.n, 'unique': row.unique_n
                }
        result = []
        for s in sends:
            d = s.to_dict()
            d['events'] = counts.get(s.id, {})
            result.append(d)
        return jsonify({'sends': result, 'total': total, 'page': page, 'per_page': per_page})

    @app.route('/campaign-history')
    @login_required
    def campaign_history():
        return render_template('campaign_history.html')

    # Social Posting                                                         #

    with app.app_context():
        db.create_all()  # creates social_tokens table if missing

    import urllib.parse as _urlparse
    import urllib.request as _urlreq
    import json as _json
    import secrets as _secrets

    def _social_cfg(key):
        return app.config.get(key, '') or ''

    @app.route('/api/social/status', methods=['GET'])
    @login_required
    def social_status():
        tokens = {t.platform: t.to_dict() for t in SocialToken.query.all()}
        return jsonify({
            'linkedin': tokens.get('linkedin', {'connected': False}),
            'facebook': tokens.get('facebook', {'connected': False}),
        })

    # ── LinkedIn OAuth ──────────────────────────────────────────────────── #

    @app.route('/social/linkedin/connect')
    @login_required
    def linkedin_connect():
        if not current_user.is_admin:
            return 'Admin only', 403
        client_id = _social_cfg('LINKEDIN_CLIENT_ID')
        if not client_id:
            return 'LINKEDIN_CLIENT_ID not set in .env', 400
        state = _secrets.token_urlsafe(16)
        from flask import session
        session['linkedin_oauth_state'] = state
        redirect_uri = _social_cfg('APP_BASE_URL') + '/social/linkedin/callback'
        params = _urlparse.urlencode({
            'response_type': 'code',
            'client_id': client_id,
            'redirect_uri': redirect_uri,
            'state': state,
            'scope': 'openid profile w_member_social',
        })
        return redirect(f'https://www.linkedin.com/oauth/v2/authorization?{params}')

    @app.route('/social/linkedin/callback')
    @login_required
    def linkedin_callback():
        from flask import session
        if not current_user.is_admin:
            return 'Admin only', 403
        error = request.args.get('error')
        if error:
            return redirect(url_for('admin_users') + '?social_error=' + error)
        code  = request.args.get('code', '')
        state = request.args.get('state', '')
        if state != session.pop('linkedin_oauth_state', None):
            return 'Invalid state', 400
        client_id     = _social_cfg('LINKEDIN_CLIENT_ID')
        client_secret = _social_cfg('LINKEDIN_CLIENT_SECRET')
        redirect_uri  = _social_cfg('APP_BASE_URL') + '/social/linkedin/callback'
        # Exchange code for token
        token_data = _urlparse.urlencode({
            'grant_type': 'authorization_code',
            'code': code,
            'redirect_uri': redirect_uri,
            'client_id': client_id,
            'client_secret': client_secret,
        }).encode()
        req = _urlreq.Request('https://www.linkedin.com/oauth/v2/accessToken',
                              data=token_data,
                              headers={'Content-Type': 'application/x-www-form-urlencoded'})
        try:
            with _urlreq.urlopen(req, timeout=10) as r:
                token_json = _json.loads(r.read())
        except Exception as exc:
            return f'Token exchange failed: {exc}', 500
        access_token = token_json.get('access_token', '')
        expires_in   = token_json.get('expires_in', 0)
        expires_at   = datetime.utcnow() + timedelta(seconds=int(expires_in)) if expires_in else None
        # Fetch profile
        profile_req = _urlreq.Request(
            'https://api.linkedin.com/v2/userinfo',
            headers={'Authorization': f'Bearer {access_token}'}
        )
        try:
            with _urlreq.urlopen(profile_req, timeout=10) as r:
                profile = _json.loads(r.read())
        except Exception:
            profile = {}
        account_name = profile.get('name') or profile.get('localizedFirstName', '') + ' ' + profile.get('localizedLastName', '')
        account_id   = profile.get('sub', '')
        tok = SocialToken.query.filter_by(platform='linkedin').first()
        if not tok:
            tok = SocialToken(platform='linkedin')
            db.session.add(tok)
        tok.access_token = access_token
        tok.account_name = account_name.strip()
        tok.account_id   = account_id
        tok.expires_at   = expires_at
        db.session.commit()
        log_audit('social_connected', 'social', None, 'linkedin')
        return redirect('/admin/users?social_connected=linkedin')

    @app.route('/social/linkedin/disconnect')
    @login_required
    def linkedin_disconnect():
        if not current_user.is_admin:
            return 'Admin only', 403
        SocialToken.query.filter_by(platform='linkedin').delete()
        db.session.commit()
        return redirect('/admin/users?social_disconnected=linkedin')

    # ── Facebook OAuth ──────────────────────────────────────────────────── #

    @app.route('/social/facebook/connect')
    @login_required
    def facebook_connect():
        if not current_user.is_admin:
            return 'Admin only', 403
        app_id = _social_cfg('FACEBOOK_APP_ID')
        if not app_id:
            return 'FACEBOOK_APP_ID not set in .env', 400
        from flask import session
        state = _secrets.token_urlsafe(16)
        session['facebook_oauth_state'] = state
        redirect_uri = _social_cfg('APP_BASE_URL') + '/social/facebook/callback'
        params = _urlparse.urlencode({
            'client_id': app_id,
            'redirect_uri': redirect_uri,
            'state': state,
            'scope': 'pages_manage_posts,pages_read_engagement,pages_show_list',
        })
        return redirect(f'https://www.facebook.com/dialog/oauth?{params}')

    @app.route('/social/facebook/callback')
    @login_required
    def facebook_callback():
        from flask import session
        if not current_user.is_admin:
            return 'Admin only', 403
        error = request.args.get('error_message') or request.args.get('error')
        if error:
            return redirect('/admin/users?social_error=' + _urlparse.quote(error))
        code  = request.args.get('code', '')
        state = request.args.get('state', '')
        if state != session.pop('facebook_oauth_state', None):
            return 'Invalid state', 400
        app_id     = _social_cfg('FACEBOOK_APP_ID')
        app_secret = _social_cfg('FACEBOOK_APP_SECRET')
        redirect_uri = _social_cfg('APP_BASE_URL') + '/social/facebook/callback'
        # Exchange code for user token
        token_url = ('https://graph.facebook.com/oauth/access_token?' +
                     _urlparse.urlencode({'client_id': app_id, 'redirect_uri': redirect_uri,
                                          'client_secret': app_secret, 'code': code}))
        try:
            with _urlreq.urlopen(token_url, timeout=10) as r:
                token_json = _json.loads(r.read())
        except Exception as exc:
            return f'Token exchange failed: {exc}', 500
        user_token = token_json.get('access_token', '')
        # Get list of pages the user manages
        pages_url = ('https://graph.facebook.com/v19.0/me/accounts?access_token=' + user_token)
        try:
            with _urlreq.urlopen(pages_url, timeout=10) as r:
                pages_json = _json.loads(r.read())
        except Exception:
            pages_json = {}
        pages = pages_json.get('data', [])
        # Use first page, or fall back to user token
        if pages:
            page      = pages[0]
            page_id   = page['id']
            page_name = page.get('name', '')
            page_token = page.get('access_token', user_token)
        else:
            page_id = page_name = ''
            page_token = user_token
        # Get user name
        me_url = 'https://graph.facebook.com/v19.0/me?access_token=' + user_token
        try:
            with _urlreq.urlopen(me_url, timeout=10) as r:
                me = _json.loads(r.read())
        except Exception:
            me = {}
        tok = SocialToken.query.filter_by(platform='facebook').first()
        if not tok:
            tok = SocialToken(platform='facebook')
            db.session.add(tok)
        tok.access_token = page_token
        tok.account_name = me.get('name', '')
        tok.page_id      = page_id
        tok.page_name    = page_name
        db.session.commit()
        log_audit('social_connected', 'social', None, 'facebook')
        return redirect('/admin/users?social_connected=facebook')

    @app.route('/social/facebook/disconnect')
    @login_required
    def facebook_disconnect():
        if not current_user.is_admin:
            return 'Admin only', 403
        SocialToken.query.filter_by(platform='facebook').delete()
        db.session.commit()
        return redirect('/admin/users?social_disconnected=facebook')

    # ── Post to Social ──────────────────────────────────────────────────── #

    @app.route('/api/social/post', methods=['POST'])
    @login_required
    def social_post():
        if not (current_user.is_admin or current_user.can_post_social):
            return jsonify({'error': 'You do not have permission to post to social media.'}), 403
        import base64 as _b64
        from flyer_render import render_flyer_png, CANVAS_FORMATS
        data       = request.get_json(force=True)
        caption    = (data.get('caption') or '').strip()
        platforms  = data.get('platforms') or []
        template_id = data.get('template_id')
        elements    = data.get('elements')
        background  = data.get('background', '#ffffff')
        bg_asset_id = data.get('bg_asset_id')

        if not caption:
            return jsonify({'error': 'Caption is required.'}), 400
        if not platforms:
            return jsonify({'error': 'Select at least one platform.'}), 400

        # Render flyer to PNG
        t = FlyerTemplate.query.get_or_404(template_id)
        els = elements or t.elements or []
        def asset_loader(asset_id):
            try:
                a = FlyerAsset.query.get(int(asset_id))
                return a.data if a else None
            except (TypeError, ValueError):
                return None
        _bg_id = bg_asset_id or t.bg_asset_id
        _bg_bytes = None
        if _bg_id:
            try:
                _ba = FlyerAsset.query.get(int(_bg_id))
                if _ba:
                    _bg_bytes = _ba.data
            except (TypeError, ValueError):
                pass
        png_bytes = render_flyer_png(els, fmt=t.format, bg_color=background, asset_loader=asset_loader, bg_image_bytes=_bg_bytes)

        results = {}

        if 'linkedin' in platforms:
            tok = SocialToken.query.filter_by(platform='linkedin').first()
            if not tok:
                results['linkedin'] = {'ok': False, 'error': 'LinkedIn not connected.'}
            else:
                try:
                    results['linkedin'] = _post_linkedin(tok.access_token, tok.account_id, caption, png_bytes)
                except Exception as exc:
                    results['linkedin'] = {'ok': False, 'error': str(exc)}

        if 'facebook' in platforms:
            tok = SocialToken.query.filter_by(platform='facebook').first()
            if not tok:
                results['facebook'] = {'ok': False, 'error': 'Facebook not connected.'}
            else:
                try:
                    results['facebook'] = _post_facebook(tok.access_token, tok.page_id, caption, png_bytes)
                except Exception as exc:
                    results['facebook'] = {'ok': False, 'error': str(exc)}

        any_ok = any(v.get('ok') for v in results.values())
        log_audit('social_post', 'social', None, None, {'platforms': platforms, 'results': results})
        return jsonify({'results': results, 'ok': any_ok})

    def _post_linkedin(access_token, author_id, caption, png_bytes):
        import base64 as _b64
        author_urn = f'urn:li:person:{author_id}'
        # Step 1: register upload
        reg_body = _json.dumps({
            'registerUploadRequest': {
                'recipes': ['urn:li:digitalmediaRecipe:feedshare-image'],
                'owner': author_urn,
                'serviceRelationships': [{'relationshipType': 'OWNER', 'identifier': 'urn:li:userGeneratedContent'}],
            }
        }).encode()
        reg_req = _urlreq.Request(
            'https://api.linkedin.com/v2/assets?action=registerUpload',
            data=reg_body,
            headers={'Authorization': f'Bearer {access_token}', 'Content-Type': 'application/json'},
        )
        with _urlreq.urlopen(reg_req, timeout=15) as r:
            reg = _json.loads(r.read())
        upload_url = reg['value']['uploadMechanism']['com.linkedin.digitalmedia.uploading.MediaUploadHttpRequest']['uploadUrl']
        asset_urn  = reg['value']['asset']
        # Step 2: upload image
        up_req = _urlreq.Request(upload_url, data=png_bytes,
                                  headers={'Authorization': f'Bearer {access_token}', 'Content-Type': 'image/png'})
        up_req.get_method = lambda: 'PUT'
        with _urlreq.urlopen(up_req, timeout=30):
            pass
        # Step 3: create post
        post_body = _json.dumps({
            'author': author_urn,
            'lifecycleState': 'PUBLISHED',
            'specificContent': {
                'com.linkedin.ugc.ShareContent': {
                    'shareCommentary': {'text': caption},
                    'shareMediaCategory': 'IMAGE',
                    'media': [{'status': 'READY', 'media': asset_urn}],
                }
            },
            'visibility': {'com.linkedin.ugc.MemberNetworkVisibility': 'PUBLIC'},
        }).encode()
        post_req = _urlreq.Request(
            'https://api.linkedin.com/v2/ugcPosts',
            data=post_body,
            headers={'Authorization': f'Bearer {access_token}', 'Content-Type': 'application/json',
                     'X-Restli-Protocol-Version': '2.0.0'},
        )
        with _urlreq.urlopen(post_req, timeout=15) as r:
            result = _json.loads(r.read())
        return {'ok': True, 'post_id': result.get('id', '')}

    def _post_facebook(access_token, page_id, caption, png_bytes):
        import io as _io
        import email.mime.multipart as _mime_mp
        import email.mime.base as _mime_base
        # POST photo to /{page-id}/photos
        boundary = _secrets.token_hex(16)
        body  = f'--{boundary}\r\nContent-Disposition: form-data; name="caption"\r\n\r\n{caption}\r\n'
        body += f'--{boundary}\r\nContent-Disposition: form-data; name="source"; filename="flyer.png"\r\nContent-Type: image/png\r\n\r\n'
        body_bytes = body.encode() + png_bytes + f'\r\n--{boundary}--\r\n'.encode()
        post_url = f'https://graph.facebook.com/v19.0/{page_id}/photos?access_token={access_token}'
        fb_req = _urlreq.Request(
            post_url, data=body_bytes,
            headers={'Content-Type': f'multipart/form-data; boundary={boundary}'},
        )
        with _urlreq.urlopen(fb_req, timeout=30) as r:
            result = _json.loads(r.read())
        return {'ok': True, 'post_id': result.get('id', '')}

    # ── Email Sequences ─────────────────────────────────────────────────── #

    @app.route('/admin/sequences')
    @login_required
    def sequences_admin():
        if not current_user.is_admin:
            return redirect('/')
        seqs = EmailSequence.query.order_by(EmailSequence.created_at.desc()).all()
        return render_template('sequences_admin.html', sequences=seqs)

    @app.route('/api/sequences', methods=['GET'])
    @login_required
    def list_sequences():
        seqs = EmailSequence.query.order_by(EmailSequence.created_at.desc()).all()
        return jsonify([s.to_dict(include_counts=True) for s in seqs])

    @app.route('/api/sequences', methods=['POST'])
    @login_required
    def create_sequence():
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        data = request.get_json(force=True)
        if not data.get('name') or not data.get('trigger_stage'):
            return jsonify({'error': 'Name and trigger stage are required'}), 400
        seq = EmailSequence(name=data['name'], trigger_stage=data['trigger_stage'])
        db.session.add(seq)
        db.session.commit()
        return jsonify(seq.to_dict()), 201

    @app.route('/api/sequences/<int:seq_id>', methods=['PATCH'])
    @login_required
    def update_sequence(seq_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        seq = EmailSequence.query.get_or_404(seq_id)
        data = request.get_json(force=True)
        if 'name' in data:
            seq.name = data['name']
        if 'trigger_stage' in data:
            seq.trigger_stage = data['trigger_stage']
        if 'is_active' in data:
            seq.is_active = bool(data['is_active'])
        db.session.commit()
        return jsonify(seq.to_dict(include_counts=True))

    @app.route('/api/sequences/<int:seq_id>', methods=['DELETE'])
    @login_required
    def delete_sequence(seq_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        seq = EmailSequence.query.get_or_404(seq_id)
        db.session.delete(seq)
        db.session.commit()
        return jsonify({'ok': True})

    @app.route('/api/sequences/<int:seq_id>/steps', methods=['POST'])
    @login_required
    def add_sequence_step(seq_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        seq = EmailSequence.query.get_or_404(seq_id)
        data = request.get_json(force=True)
        step = EmailSequenceStep(
            sequence_id=seq_id,
            step_order=len(seq.steps),
            day_offset=int(data.get('day_offset', 1)),
            subject=data.get('subject', ''),
            body=data.get('body', ''),
        )
        db.session.add(step)
        db.session.commit()
        return jsonify(step.to_dict()), 201

    @app.route('/api/sequences/steps/<int:step_id>', methods=['PATCH'])
    @login_required
    def update_sequence_step(step_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        step = EmailSequenceStep.query.get_or_404(step_id)
        data = request.get_json(force=True)
        for field in ['day_offset', 'subject', 'body']:
            if field in data:
                setattr(step, field, data[field])
        db.session.commit()
        return jsonify(step.to_dict())

    @app.route('/api/sequences/steps/<int:step_id>', methods=['DELETE'])
    @login_required
    def delete_sequence_step(step_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        step = EmailSequenceStep.query.get_or_404(step_id)
        db.session.delete(step)
        db.session.commit()
        return jsonify({'ok': True})

    @app.route('/api/sequences/<int:seq_id>/process-now', methods=['POST'])
    @login_required
    def process_sequence_now(seq_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        _process_sequence_emails(app)
        return jsonify({'ok': True})

    # ── Landing Pages ───────────────────────────────────────────────────── #

    @app.route('/admin/landing-pages')
    @login_required
    def landing_pages_admin():
        if not current_user.is_admin:
            return redirect('/')
        pages = LandingPage.query.order_by(LandingPage.created_at.desc()).all()
        return render_template('landing_pages.html', pages=pages)

    @app.route('/admin/landing-pages/<int:page_id>/edit')
    @login_required
    def landing_page_edit(page_id):
        if not current_user.is_admin:
            return redirect('/')
        page = LandingPage.query.get_or_404(page_id)
        return render_template('landing_page_editor.html', page=page)

    @app.route('/api/landing-pages', methods=['POST'])
    @login_required
    def create_landing_page():
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        data = request.get_json(force=True)
        if not data.get('title'):
            return jsonify({'error': 'Title is required'}), 400
        import re, uuid
        raw = data.get('slug') or re.sub(r'[^a-z0-9]+', '-', data['title'].lower()).strip('-')
        slug = raw[:80]
        if LandingPage.query.filter_by(slug=slug).first():
            slug = slug[:74] + '-' + uuid.uuid4().hex[:5]
        page = LandingPage(
            slug=slug,
            title=data['title'],
            subtitle=data.get('subtitle') or None,
            body=data.get('body') or None,
            bg_color=data.get('bg_color', '#ffffff'),
            text_color=data.get('text_color', '#111111'),
            button_text=data.get('button_text', 'Get in Touch'),
            button_color=data.get('button_color', '#AD0304'),
            show_phone=bool(data.get('show_phone', True)),
            show_message=bool(data.get('show_message', True)),
            pipeline_stage=data.get('pipeline_stage') or None,
        )
        db.session.add(page)
        db.session.commit()
        return jsonify(page.to_dict()), 201

    @app.route('/api/landing-pages/<int:page_id>', methods=['PATCH'])
    @login_required
    def update_landing_page(page_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        page = LandingPage.query.get_or_404(page_id)
        data = request.get_json(force=True)
        for field in ['title', 'subtitle', 'body', 'bg_color', 'text_color',
                      'button_text', 'button_color', 'pipeline_stage']:
            if field in data:
                setattr(page, field, data[field] or None if field in ('subtitle', 'body', 'pipeline_stage') else data[field])
        for field in ['show_phone', 'show_message', 'is_active']:
            if field in data:
                setattr(page, field, bool(data[field]))
        db.session.commit()
        return jsonify(page.to_dict())

    @app.route('/api/landing-pages/<int:page_id>', methods=['DELETE'])
    @login_required
    def delete_landing_page(page_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        page = LandingPage.query.get_or_404(page_id)
        db.session.delete(page)
        db.session.commit()
        return jsonify({'ok': True})

    @app.route('/api/landing-pages/<int:page_id>/submissions')
    @login_required
    def get_lp_submissions(page_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        subs = LandingPageSubmission.query.filter_by(page_id=page_id)\
            .order_by(LandingPageSubmission.created_at.desc()).all()
        return jsonify([s.to_dict() for s in subs])

    @app.route('/p/<slug>')
    def public_landing_page(slug):
        page = LandingPage.query.filter_by(slug=slug, is_active=True).first_or_404()
        return render_template('landing_page_public.html', page=page)

    @app.route('/p/<slug>/submit', methods=['POST'])
    def submit_landing_page(slug):
        page = LandingPage.query.filter_by(slug=slug, is_active=True).first_or_404()
        data = request.get_json(force=True)
        if not data.get('name') or not data.get('email'):
            return jsonify({'error': 'Name and email are required'}), 400

        email = data['email'].strip().lower()
        name  = data['name'].strip()
        parts = name.rsplit(' ', 1)
        first = parts[0]
        last  = parts[1] if len(parts) > 1 else ''

        # Match or create contact
        contact = Contact.query.filter(
            db.func.lower(Contact.email) == email
        ).first()
        if not contact:
            contact = Contact(
                first_name=first,
                last_name=last,
                email=email,
                phone=(data.get('phone') or '').strip() or None,
                pipeline_stage=page.pipeline_stage or None,
            )
            db.session.add(contact)
        else:
            if page.pipeline_stage and not contact.pipeline_stage:
                contact.pipeline_stage = page.pipeline_stage

        db.session.flush()

        activity = Activity(
            contact_id=contact.id,
            activity_type='note',
            notes=f'Submitted landing page "{page.title}"' +
                  (f': {data["message"]}' if data.get('message') else ''),
            created_by=None,
        )
        db.session.add(activity)

        sub = LandingPageSubmission(
            page_id=page.id,
            name=name,
            email=email,
            phone=(data.get('phone') or '').strip() or None,
            message=(data.get('message') or '').strip() or None,
            contact_id=contact.id,
        )
        db.session.add(sub)
        db.session.commit()
        return jsonify({'ok': True}), 201

    # ── Meeting Scheduler ───────────────────────────────────────────────── #

    @app.route('/book')
    def public_book():
        return render_template('book.html')

    @app.route('/admin/scheduler')
    @login_required
    def scheduler_admin():
        if not current_user.is_admin:
            return redirect('/')
        rules = AvailabilityRule.query.order_by(AvailabilityRule.day_of_week).all()
        return render_template('scheduler_admin.html', rules=rules)

    @app.route('/api/scheduler/availability', methods=['GET'])
    @login_required
    def get_availability():
        rules = AvailabilityRule.query.order_by(AvailabilityRule.day_of_week).all()
        return jsonify([r.to_dict() for r in rules])

    @app.route('/api/scheduler/availability', methods=['POST'])
    @login_required
    def save_availability():
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        data = request.get_json(force=True)
        # data = list of {day_of_week, start_hour, end_hour, slot_minutes, enabled}
        AvailabilityRule.query.delete()
        for item in data:
            if item.get('enabled'):
                db.session.add(AvailabilityRule(
                    day_of_week=int(item['day_of_week']),
                    start_hour=int(item['start_hour']),
                    end_hour=int(item['end_hour']),
                    slot_minutes=int(item.get('slot_minutes', 30)),
                ))
        db.session.commit()
        return jsonify({'ok': True})

    @app.route('/api/scheduler/slots')
    def get_slots():
        from datetime import date as _date, time as _time, timedelta
        date_str = request.args.get('date', '')
        try:
            req_date = _date.fromisoformat(date_str)
        except (ValueError, TypeError):
            return jsonify({'error': 'Invalid date'}), 400

        dow = req_date.weekday()  # 0=Mon
        rule = AvailabilityRule.query.filter_by(day_of_week=dow).first()
        if not rule:
            return jsonify([])

        # Existing confirmed bookings for that date
        booked_starts = {
            b.start_time.strftime('%H:%M')
            for b in Booking.query.filter_by(date=req_date, status='confirmed').all()
        }

        slots = []
        slot_delta = timedelta(minutes=rule.slot_minutes)
        from datetime import datetime as _dt
        current = _dt.combine(req_date, _time(rule.start_hour, 0))
        end_dt  = _dt.combine(req_date, _time(rule.end_hour, 0))
        while current + slot_delta <= end_dt:
            label = current.strftime('%H:%M')
            if label not in booked_starts:
                slots.append(label)
            current += slot_delta

        return jsonify(slots)

    @app.route('/api/scheduler/book', methods=['POST'])
    def create_booking():
        from datetime import date as _date, time as _time, timedelta
        data = request.get_json(force=True)
        required = ['date', 'start_time', 'name', 'email']
        missing = [f for f in required if not data.get(f)]
        if missing:
            return jsonify({'error': f'Missing: {", ".join(missing)}'}), 400

        try:
            req_date = _date.fromisoformat(data['date'])
        except ValueError:
            return jsonify({'error': 'Invalid date'}), 400

        dow = req_date.weekday()
        rule = AvailabilityRule.query.filter_by(day_of_week=dow).first()
        if not rule:
            return jsonify({'error': 'No availability on that day'}), 409

        try:
            h, m = map(int, data['start_time'].split(':'))
            start = _time(h, m)
        except Exception:
            return jsonify({'error': 'Invalid time'}), 400

        from datetime import datetime as _dt
        end_dt = (_dt.combine(req_date, start) + timedelta(minutes=rule.slot_minutes)).time()

        # Check not already booked
        conflict = Booking.query.filter_by(
            date=req_date, start_time=start, status='confirmed'
        ).first()
        if conflict:
            return jsonify({'error': 'That slot is no longer available'}), 409

        booking = Booking(
            date=req_date,
            start_time=start,
            end_time=end_dt,
            name=data['name'].strip(),
            email=data['email'].strip().lower(),
            phone=(data.get('phone') or '').strip() or None,
            notes=(data.get('notes') or '').strip() or None,
        )
        db.session.add(booking)
        db.session.commit()

        # Send confirmation email
        try:
            sg_key = os.environ.get('SENDGRID_API_KEY') or (
                os.environ.get('SMTP_PASSWORD', '').startswith('SG.') and os.environ.get('SMTP_PASSWORD')
            ) or None
            from_email = os.environ.get('MAIL_FROM', os.environ.get('SMTP_USERNAME', ''))
            if sg_key and from_email:
                day_label = req_date.strftime('%A, %B %-d, %Y')
                time_label = _dt.combine(req_date, start).strftime('%-I:%M %p')
                html = (
                    f'<p>Hi {booking.name},</p>'
                    f'<p>Your meeting with JBJ Management has been confirmed for '
                    f'<strong>{day_label} at {time_label}</strong>.</p>'
                    f'<p>If you need to cancel or reschedule, please reply to this email.</p>'
                    f'<p>- JBJ Management</p>'
                )
                import json as _json2
                import urllib.request as _urlreq2
                payload = {
                    'personalizations': [{'to': [{'email': booking.email, 'name': booking.name}]}],
                    'from': {'email': from_email},
                    'subject': f'Meeting confirmed - {day_label} at {time_label}',
                    'content': [{'type': 'text/html', 'value': html}],
                }
                req = _urlreq2.Request(
                    'https://api.sendgrid.com/v3/mail/send',
                    data=_json2.dumps(payload).encode(),
                    headers={
                        'Authorization': f'Bearer {sg_key}',
                        'Content-Type': 'application/json',
                    },
                    method='POST',
                )
                _urlreq2.urlopen(req, timeout=10)
        except Exception:
            pass  # Booking is saved; email failure is non-fatal

        return jsonify({'ok': True, 'booking': booking.to_dict()}), 201

    @app.route('/api/scheduler/bookings')
    @login_required
    def list_bookings():
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        from datetime import date as _date
        upcoming = Booking.query.filter(
            Booking.date >= _date.today()
        ).order_by(Booking.date, Booking.start_time).all()
        past = Booking.query.filter(
            Booking.date < _date.today()
        ).order_by(Booking.date.desc(), Booking.start_time).limit(50).all()
        return jsonify({'upcoming': [b.to_dict() for b in upcoming],
                        'past': [b.to_dict() for b in past]})

    @app.route('/api/scheduler/bookings/<int:booking_id>', methods=['PATCH'])
    @login_required
    def update_booking(booking_id):
        if not current_user.is_admin:
            return jsonify({'error': 'Admins only'}), 403
        booking = Booking.query.get_or_404(booking_id)
        data = request.get_json(force=True)
        if 'status' in data:
            booking.status = data['status']
        db.session.commit()
        return jsonify(booking.to_dict())

    return app


if __name__ == '__main__':
    app = create_app()
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), threaded=True)
